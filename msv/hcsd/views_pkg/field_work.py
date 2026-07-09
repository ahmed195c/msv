"""
Field Work Orders views.

URL prefix : /field-work/
Templates  : hcsd/field_work_*.html
"""

import datetime as _dt
import io
import logging
import os

from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse
from django.utils import timezone

from ..models import FieldWorkOrder, FieldWorkOrderLog, FieldWorkPhoto, FieldWorkSupervisorArea, FieldWorkSupervisorProfile
from .common import _can_admin, _can_data_entry, _can_fw_supervise, _fw_supervisor_users_qs

logger = logging.getLogger(__name__)

_STATUS_LABELS = dict(FieldWorkOrder.STATUS_CHOICES)


def _fw_log(order, action, actor, from_value='', to_value='', note=''):
    FieldWorkOrderLog.objects.create(
        order=order, action=action, actor=actor,
        from_value=from_value, to_value=to_value, note=note,
    )


ALLOWED_PHOTO_EXTENSIONS = {'.jpg', '.jpeg', '.png'}

_FW_CLOSED_STATUSES = frozenset({
    'completed',
    'private_company', 'cust_declined', 'wrong_phone', 'phone_off', 'no_answer',
    'other_municipal',
    'closed_private_building', 'closed_no_answer', 'closed_other_municipal',
    'closed_observation', 'closed_low_infestation', 'closed_moderate_infestation',
    'closed_high_infestation', 'closed_out_of_service', 'closed_customer_refused',
    'closed_mobile_off', 'closed_not_attending', 'closed_not_available',
    'closed_scheduled_client',
})
# Closed without completion (refused, no answer, etc.) — excludes 'completed'
_FW_TRULY_CLOSED = _FW_CLOSED_STATUSES - {'completed'}


def _infer_status_from_excel(excel_status: str) -> str:
    """Map raw Excel status text to a system status code."""
    s = excel_status.upper().strip()
    if not s:
        return 'new'
    if any(k in s for k in ('COMPLET', 'DONE', 'FINISHED', 'ACCOMPLISHED', 'SERVICE HAS BEEN')):
        return 'completed'
    if any(k in s for k in ('NO ANSWER', 'NOT ANSWER', 'NO RESPONSE', 'NOT RESPOND', 'NOT REPLYING')):
        return 'closed_no_answer'
    if any(k in s for k in ('REFUSED', 'DECLINED', 'REJECTED', 'REFUSED SERVICE')):
        return 'closed_customer_refused'
    if any(k in s for k in ('PHONE OFF', 'MOBILE OFF', 'SWITCHED OFF')):
        return 'closed_mobile_off'
    if any(k in s for k in ('WRONG PHONE', 'WRONG NUMBER', 'INVALID PHONE')):
        return 'wrong_phone'
    if any(k in s for k in ('POSTPONED', 'RESCHEDULED', 'DELAYED')):
        return 'postponed_client'
    if any(k in s for k in ('PRIVATE COMPANY', 'PRIVATE BUILDING', 'PRIVATE')):
        return 'closed_private_building'
    if any(k in s for k in ('OTHER MUNICIPAL', 'ANOTHER MUNICIPAL', 'OTHER AUTHORITY')):
        return 'closed_other_municipal'
    if any(k in s for k in ('CLOSED', 'CLOSE')):
        return 'closed_observation'
    if any(k in s for k in ('NOT AVAILABLE', 'UNAVAILABLE')):
        return 'closed_not_available'
    if any(k in s for k in ('NOT ATTENDING', 'NOT ATTENDING')):
        return 'closed_not_attending'
    return 'new'


# ---------------------------------------------------------------------------
# List
# ---------------------------------------------------------------------------

@login_required
def field_work_list(request):
    from django.core.paginator import Paginator
    from django.db.models import Q, Case, When, IntegerField, ExpressionWrapper, DurationField, F, Count
    from django.db.models.functions import Now

    can_admin      = _can_admin(request.user)
    can_data_entry = _can_data_entry(request.user)

    # ── Bulk delete POST ─────────────────────────────────────────────────────
    if request.method == 'POST' and request.POST.get('action') == 'bulk_delete':
        if not can_admin:
            from django.http import HttpResponseForbidden
            return HttpResponseForbidden()
        order_ids = [int(i) for i in request.POST.getlist('order_ids') if i.isdigit()]
        if order_ids:
            FieldWorkOrder.objects.filter(pk__in=order_ids).delete()
        from django.shortcuts import redirect
        qs = request.META.get('QUERY_STRING', '')
        return redirect(request.path + ('?' + qs if qs else ''))

    # ── Bulk assign POST ─────────────────────────────────────────────────────
    if request.method == 'POST' and request.POST.get('action') == 'bulk_assign':
        if not (can_admin or can_data_entry):
            from django.http import HttpResponseForbidden
            return HttpResponseForbidden()
        sup_id   = (request.POST.get('supervisor_id') or '').strip()
        order_ids = request.POST.getlist('order_ids')
        try:
            sup_user  = _fw_supervisor_users_qs().select_related('fw_supervisor_profile').get(pk=int(sup_id))
            sup_label = sup_user.get_full_name() or sup_user.username
            bulk_orders = FieldWorkOrder.objects.filter(pk__in=[int(i) for i in order_ids if i.isdigit()])
            for order in bulk_orders:
                old_sup   = order.assigned_supervisor
                old_label = (old_sup.get_full_name() or old_sup.username) if old_sup else ''
                log_action = 'reassigned' if old_sup else 'assigned'
                order.assigned_supervisor = sup_user
                order.assigned_at = timezone.now()
                fields = ['assigned_supervisor', 'assigned_at']
                if order.status in ('new', 'pending', 'postponed_client'):
                    order.status = 'supervisor_assigned'
                    order.postponed_until = None
                    fields += ['status', 'postponed_until']
                order.save(update_fields=fields)
                _fw_log(order, log_action, request.user, from_value=old_label, to_value=sup_label)
        except (ValueError, Exception):
            pass
        from django.shortcuts import redirect
        return redirect(request.get_full_path().replace('?' , '?').split('?')[0] +
                        ('?' + request.META.get('QUERY_STRING', '') if request.META.get('QUERY_STRING') else ''))
    # ────────────────────────────────────────────────────────────────────────

    status_filter   = (request.GET.get('status')      or 'all').strip()
    source_filter   = (request.GET.get('source')      or 'all').strip()
    search          = (request.GET.get('q')            or '').strip()
    quick_filter    = (request.GET.get('quick')        or '').strip()
    date_from_req   = (request.GET.get('date_from')   or '').strip()
    date_to_req     = (request.GET.get('date_to')     or '').strip()
    date_type       = (request.GET.get('date_type')   or 'request').strip()  # 'request' or 'close'

    import datetime as _datetime_mod
    _date_from = None
    _date_to   = None
    try:
        _date_from = _datetime_mod.date.fromisoformat(date_from_req)
    except ValueError:
        pass
    try:
        _date_to = _datetime_mod.date.fromisoformat(date_to_req)
    except ValueError:
        pass

    _CLOSED_STATUSES = [s for s, _ in FieldWorkOrder.STATUS_CHOICES if s.startswith('closed_')]

    orders = FieldWorkOrder.objects.only(
        'id', 'order_number', 'customer_name', 'site_name',
        'area', 'location', 'pest_types', 'request_date', 'work_date',
        'close_date', 'postponed_until', 'supervisor_name', 'status', 'excel_status', 'source',
        'created_at', 'assigned_supervisor',
    ).select_related('assigned_supervisor', 'assigned_supervisor__fw_supervisor_profile')

    # Supervisors see only their area orders + directly assigned/received
    if _can_fw_supervise(request.user) and not _can_admin(request.user) and not _can_data_entry(request.user):
        my_areas = list(
            FieldWorkSupervisorArea.objects.filter(supervisor=request.user)
            .values_list('area', flat=True)
        )
        orders = orders.filter(
            Q(area__in=my_areas)
            | Q(assigned_supervisor=request.user)
            | Q(received_by=request.user)
        ).distinct()

    # Quick filter overrides the status dropdown
    if quick_filter == 'today':
        _today = timezone.localdate()
        orders = orders.filter(Q(request_date=_today) | Q(request_date__isnull=True, created_at__date=_today))
    elif quick_filter == 'new':
        _today = timezone.localdate()
        orders = orders.filter(
            Q(status='new') |
            Q(status='postponed_client', postponed_until__lte=_today)
        )
    elif quick_filter == 'received':
        orders = orders.filter(status__in=['supervisor_assigned', 'order_received'])
    elif quick_filter == 'completed':
        orders = orders.filter(status='completed')
    elif quick_filter == 'closed':
        orders = orders.filter(status__in=_FW_TRULY_CLOSED)
    elif quick_filter == 'postponed':
        orders = orders.filter(status='postponed_client')
    elif status_filter == 'closed':
        orders = orders.filter(status__in=_FW_TRULY_CLOSED)
    elif status_filter != 'all':
        orders = orders.filter(status=status_filter)
    else:
        # Orders with an assigned supervisor only appear under the "طلبات مسلَّمة" quick filter
        orders = orders.exclude(status__in=['supervisor_assigned', 'order_received'])

    if source_filter != 'all':
        orders = orders.filter(source=source_filter)
    if search:
        orders = orders.filter(
            Q(order_number__icontains=search)
            | Q(customer_name__icontains=search)
            | Q(mobile__icontains=search)
            | Q(area__icontains=search)
            | Q(supervisor_name__icontains=search)
            | Q(work_type__icontains=search)
            | Q(site_name__icontains=search)
            | Q(assigned_supervisor__fw_supervisor_profile__admin_number__icontains=search)
            | Q(assigned_supervisor__fw_supervisor_profile__name_ar__icontains=search)
            | Q(assigned_supervisor__fw_supervisor_profile__name_en__icontains=search)
        )

    # Date range filter
    if _date_from or _date_to:
        if date_type == 'close':
            if _date_from:
                orders = orders.filter(close_date__gte=_date_from)
            if _date_to:
                orders = orders.filter(close_date__lte=_date_to)
        else:
            if _date_from:
                orders = orders.filter(
                    Q(request_date__gte=_date_from) |
                    Q(request_date__isnull=True, created_at__date__gte=_date_from)
                )
            if _date_to:
                orders = orders.filter(
                    Q(request_date__lte=_date_to) |
                    Q(request_date__isnull=True, created_at__date__lte=_date_to)
                )

    sort = (request.GET.get('sort') or '').strip()
    sort_dir = (request.GET.get('sort_dir') or 'desc').strip()
    if sort_dir not in ('asc', 'desc'):
        sort_dir = 'desc'
    _desc = sort_dir == 'desc'

    if sort == 'area':
        orders = orders.order_by('area' if _desc else '-area', '-created_at', '-pk')
    elif sort == 'customer':
        from django.db.models import Subquery, OuterRef
        cust_freq = (
            FieldWorkOrder.objects.filter(customer_name=OuterRef('customer_name'))
            .values('customer_name')
            .annotate(_n=Count('id'))
            .values('_n')
        )
        orders = orders.annotate(_cust_freq=Subquery(cust_freq, output_field=IntegerField()))
        _freq_order = '-_cust_freq' if _desc else '_cust_freq'
        orders = orders.order_by(_freq_order, 'customer_name', '-created_at', '-pk')
    elif sort == 'supervisor':
        from django.db.models import Subquery, OuterRef
        sup_freq = (
            FieldWorkOrder.objects.filter(supervisor_name=OuterRef('supervisor_name'))
            .values('supervisor_name')
            .annotate(_n=Count('id'))
            .values('_n')
        )
        orders = orders.annotate(_sup_freq=Subquery(sup_freq, output_field=IntegerField()))
        _freq_order = '-_sup_freq' if _desc else '_sup_freq'
        orders = orders.order_by(_freq_order, 'supervisor_name', '-created_at', '-pk')
    else:
        # Default ordering: due-postponed first → new → received → completed/closed
        _today_ord = timezone.localdate()
        orders = orders.annotate(
            _priority=Case(
                When(status='postponed_client', postponed_until__lte=_today_ord, then=0),
                When(status__in=['new', 'supervisor_assigned'], then=1),
                When(status='order_received', then=2),
                When(status='completed', then=3),
                default=4,
                output_field=IntegerField(),
            )
        ).order_by('_priority', 'postponed_until', '-created_at', '-pk')

    # Build sort URLs: clicking active column toggles direction; clicking new column resets to desc
    _qs = request.GET.copy()
    _qs.pop('page', None)
    _sort_urls = {}
    for _s in ('area', 'customer', 'supervisor'):
        _qs['sort'] = _s
        _qs['sort_dir'] = 'asc' if (_s == sort and _desc) else 'desc'
        _sort_urls[_s] = '?' + _qs.urlencode()
    _qs.pop('sort', None)
    _qs.pop('sort_dir', None)
    _sort_urls['default'] = ('?' + _qs.urlencode()) if _qs else '?'

    paginator = Paginator(orders, 50)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    # Compute days_old and due-postponed flag in Python
    _today = timezone.localdate()
    for _order in page_obj.object_list:
        _order.days_old_int = (_today - _order.created_at.date()).days
        _order.is_due_postponed = (
            _order.status == 'postponed_client' and
            bool(_order.postponed_until) and
            _order.postponed_until <= _today
        )

    _active_choices = [(v, l) for v, l in FieldWorkOrder.STATUS_CHOICES if v not in _FW_TRULY_CLOSED]
    status_options = [('all', 'كل الحالات — All')] + _active_choices + [('closed', 'مغلق — Closed')]

    can_assign = can_admin or can_data_entry
    fw_supervisors = _fw_supervisor_users_qs().prefetch_related('fw_supervisor_profile') if can_assign else []

    return render(request, 'hcsd/field_work_list.html', {
        'page_obj':        page_obj,
        'page_start':      page_obj.start_index,
        'can_admin':       can_admin,
        'can_data_entry':  can_data_entry,
        'can_assign':      can_assign,
        'fw_supervisors':  fw_supervisors,
        'status_filter':   status_filter,
        'source_filter':   source_filter,
        'search':          search,
        'quick_filter':    quick_filter,
        'today':           _today,
        'status_options':  status_options,
        'total_count':     paginator.count,
        'current_sort':    sort,
        'sort_dir':        sort_dir,
        'sort_urls':       _sort_urls,
        'date_from_req':   date_from_req,
        'date_to_req':     date_to_req,
        'date_type':       date_type,
    })


# ---------------------------------------------------------------------------
# Create
# ---------------------------------------------------------------------------

@login_required
def field_work_create(request):
    if not (_can_admin(request.user) or _can_data_entry(request.user)):
        return redirect('field_work_list')

    errors = []

    if request.method == 'POST':
        import datetime as _dt_cls
        get = lambda k: (request.POST.get(k) or '').strip()

        order_number      = get('order_number')
        customer_name     = get('customer_name')
        mobile            = get('mobile')
        area              = get('area')
        street_number     = get('street_number')
        house_number      = get('house_number')
        site_name         = get('site_name')
        location          = get('location')
        pest_types        = get('pest_types')
        complaint_source  = get('complaint_source')
        status            = get('status')
        notes             = get('notes')
        request_date_raw  = get('request_date')

        _valid_sources = {c for c, _ in FieldWorkOrder.COMPLAINT_SOURCE_CHOICES}
        _valid_statuses = {c for c, _ in FieldWorkOrder.STATUS_CHOICES}

        if not customer_name and not order_number:
            errors.append('يرجى إدخال اسم المتعامل أو رقم الطلب على الأقل.')
        if complaint_source and complaint_source not in _valid_sources:
            complaint_source = ''
        if status not in _valid_statuses:
            status = 'new'

        def _parse_date(s):
            if not s:
                return None
            try:
                return _dt_cls.date.fromisoformat(s)
            except ValueError:
                return None

        if not errors:
            order = FieldWorkOrder.objects.create(
                order_number=order_number,
                customer_name=customer_name,
                mobile=mobile,
                area=area,
                street_number=street_number,
                house_number=house_number,
                site_name=site_name,
                location=location,
                pest_types=pest_types,
                complaint_source=complaint_source,
                status=status,
                notes=notes,
                request_date=_parse_date(request_date_raw),
                source='manual',
                created_by=request.user,
            )
            return redirect('field_work_detail', pk=order.pk)

    return render(request, 'hcsd/field_work_create.html', {
        'errors':             errors,
        'post':               request.POST,
        'status_choices':     FieldWorkOrder.STATUS_CHOICES,
        'complaint_sources':  FieldWorkOrder.COMPLAINT_SOURCE_CHOICES,
    })


# ---------------------------------------------------------------------------
# Detail
# ---------------------------------------------------------------------------

@login_required
def field_work_detail(request, pk):
    order = get_object_or_404(
        FieldWorkOrder.objects.select_related(
            'created_by', 'assigned_supervisor', 'received_by',
            'report_submitted_by', 'report_submitted_by__fw_supervisor_profile',
            'location_saved_by',
        ).prefetch_related('photos', 'logs__actor'), pk=pk
    )
    can_admin = _can_admin(request.user)
    can_data_entry = _can_data_entry(request.user)
    can_edit = can_admin or can_data_entry
    is_fw_supervisor = _can_fw_supervise(request.user)
    can_assign = can_admin or can_data_entry
    uid = request.user.id
    can_submit_report = can_admin or (
        is_fw_supervisor and (
            order.assigned_supervisor_id == uid or order.received_by_id == uid
        )
    )
    can_receive = (
        is_fw_supervisor
        and not can_submit_report
        and order.status not in _FW_CLOSED_STATUSES
    )

    photos_all = sorted(order.photos.all(), key=lambda p: p.uploaded_at or p.pk)

    errors = []
    success = None

    if request.method == 'POST':
        action = (request.POST.get('action') or '').strip()

        # ── Receive order ────────────────────────────────────────────────────
        if action == 'receive_order' and is_fw_supervisor and order.status != 'completed':
            order.received_by = request.user
            order.received_at = timezone.now()
            save_fields = ['received_by', 'received_at']
            if order.status not in ('completed',):
                order.status = 'order_received'
                save_fields.append('status')
            order.save(update_fields=save_fields)
            _fw_log(order, 'received', request.user)
            success = 'تم استلام متابعة الأمر.'

        # ── Assign supervisor ────────────────────────────────────────────────
        elif action == 'assign_supervisor' and can_assign:
            sup_id = (request.POST.get('supervisor_id') or '').strip()
            old_sup = order.assigned_supervisor
            old_label = old_sup.get_full_name() or old_sup.username if old_sup else ''
            if sup_id == '':
                order.assigned_supervisor = None
                order.assigned_at = None
                save_fields = ['assigned_supervisor', 'assigned_at']
                if order.status in ('supervisor_assigned', 'pending'):
                    order.status = 'new'
                    save_fields.append('status')
                order.save(update_fields=save_fields)
                _fw_log(order, 'unassigned', request.user, from_value=old_label)
                success = 'تم إلغاء تعيين المراقب.'
            else:
                try:
                    sup_user = _fw_supervisor_users_qs().get(pk=int(sup_id))
                    log_action = 'reassigned' if old_sup else 'assigned'
                    order.assigned_supervisor = sup_user
                    order.assigned_at = timezone.now()
                    save_fields = ['assigned_supervisor', 'assigned_at']
                    if order.status in ('new', 'pending', 'postponed_client'):
                        order.status = 'supervisor_assigned'
                        order.postponed_until = None
                        save_fields += ['status', 'postponed_until']
                    order.save(update_fields=save_fields)
                    new_label = sup_user.get_full_name() or sup_user.username
                    _fw_log(order, log_action, request.user, from_value=old_label, to_value=new_label)
                    success = f'تم تعيين {new_label} مراقباً للأمر.'
                except (ValueError, Exception):
                    errors.append('المراقب المختار غير صالح.')

        elif not can_edit and not can_submit_report:
            errors.append('ليس لديك صلاحية لتنفيذ هذا الإجراء.')

        # ── Update work details ──────────────────────────────────────────────
        elif action == 'update_details' and can_edit:
            complaint_source_upd = (request.POST.get('complaint_source') or '').strip()
            _valid_src = {c for c, _ in FieldWorkOrder.COMPLAINT_SOURCE_CHOICES}
            if complaint_source_upd not in _valid_src:
                complaint_source_upd = ''
            site_name   = (request.POST.get('site_name') or '').strip()
            work_type   = (request.POST.get('work_type') or '').strip()
            location    = (request.POST.get('location') or '').strip()
            description = (request.POST.get('description') or '').strip()
            work_date   = (request.POST.get('work_date') or '').strip() or None
            notes       = (request.POST.get('notes') or '').strip()
            workers_count_raw = (request.POST.get('workers_count') or '').strip()
            equipment   = (request.POST.get('equipment_used') or '').strip()
            # Excel-specific fields
            customer_name = (request.POST.get('customer_name') or '').strip()
            mobile        = (request.POST.get('mobile') or '').strip()
            area          = (request.POST.get('area') or '').strip()
            street_number = (request.POST.get('street_number') or '').strip()
            house_number  = (request.POST.get('house_number') or '').strip()
            pest_types    = (request.POST.get('pest_types') or '').strip()
            close_date_raw = (request.POST.get('close_date') or '').strip() or None

            workers_count = None
            if workers_count_raw:
                try:
                    workers_count = int(workers_count_raw)
                    if workers_count < 0:
                        raise ValueError
                except ValueError:
                    errors.append('عدد العمال يجب أن يكون رقماً صحيحاً.')

            if not errors:
                order.complaint_source = complaint_source_upd
                order.site_name = site_name
                order.work_type = work_type
                order.location = location
                order.description = description
                order.work_date = work_date
                order.notes = notes
                order.workers_count = workers_count
                order.equipment_used = equipment
                order.customer_name = customer_name
                order.mobile = mobile
                order.area = area
                order.street_number = street_number
                order.house_number = house_number
                order.pest_types = pest_types
                if close_date_raw:
                    try:
                        from datetime import date as _date_cls
                        order.close_date = _date_cls.fromisoformat(close_date_raw)
                    except ValueError:
                        pass
                order.save(update_fields=[
                    'complaint_source',
                    'site_name', 'work_type', 'location', 'description',
                    'work_date', 'notes', 'workers_count', 'equipment_used',
                    'customer_name', 'mobile', 'area', 'street_number',
                    'house_number', 'pest_types', 'close_date',
                ])
                success = 'تم حفظ التفاصيل.'

        # ── Update status ────────────────────────────────────────────────────
        elif action == 'update_status' and can_edit:
            new_status = (request.POST.get('status') or '').strip()
            work_completed_raw = request.POST.get('work_completed')
            valid_statuses = {s for s, _ in FieldWorkOrder.STATUS_CHOICES}
            if new_status not in valid_statuses:
                errors.append('حالة غير صحيحة.')
            else:
                old_status = order.status
                order.status = new_status
                if new_status in ('completed', 'incomplete'):
                    order.work_completed = (new_status == 'completed')
                order.save(update_fields=['status', 'work_completed'])
                _fw_log(order, 'status_changed', request.user,
                        from_value=_STATUS_LABELS.get(old_status, old_status),
                        to_value=_STATUS_LABELS.get(new_status, new_status))
                success = 'تم تحديث الحالة.'

        # ── Upload photos ────────────────────────────────────────────────────
        elif action == 'upload_photos' and can_edit:
            photos = request.FILES.getlist('photos')
            if not photos:
                errors.append('يرجى اختيار صورة واحدة على الأقل.')
            else:
                invalid = [
                    p.name for p in photos
                    if os.path.splitext(p.name)[1].lower() not in ALLOWED_PHOTO_EXTENSIONS
                ]
                if invalid:
                    errors.append('يُسمح فقط بصور JPG/PNG.')
                else:
                    for photo in photos:
                        FieldWorkPhoto.objects.create(
                            work_order=order,
                            phase='work',
                            file=photo,
                            uploaded_by=request.user,
                        )
                    success = 'تم رفع الصور.'
                    photos_all = order.photos.order_by('uploaded_at')

        # ── Save GPS location ────────────────────────────────────────────────
        elif action == 'save_location' and can_submit_report:
            try:
                lat = float(request.POST.get('lat', ''))
                lng = float(request.POST.get('lng', ''))
            except (ValueError, TypeError):
                errors.append('إحداثيات غير صالحة.')
            else:
                order.gps_lat = lat
                order.gps_lng = lng
                order.location_saved_at = timezone.now()
                order.location_saved_by = request.user
                update_f = ['gps_lat', 'gps_lng', 'location_saved_at', 'location_saved_by']
                if not order.time_in:
                    order.time_in = timezone.now()
                    update_f.append('time_in')
                order.save(update_fields=update_f)
                success = 'تم حفظ الموقع.'

        # ── Supervisor report ────────────────────────────────────────────────
        elif action == 'supervisor_report' and can_submit_report:
            import json as _json
            workers_raw    = (request.POST.get('workers_count') or '').strip()
            vehicles_raw   = (request.POST.get('vehicles_count') or '').strip()
            building_type  = (request.POST.get('building_type') or '').strip()
            sup_notes      = (request.POST.get('supervisor_notes') or '').strip()
            client_sig     = (request.POST.get('client_signature') or '').strip()
            supervisor_sig = (request.POST.get('supervisor_signature') or '').strip()

            try:
                spray_entries = _json.loads(request.POST.get('spray_entries_json') or '[]')
                if not isinstance(spray_entries, list):
                    spray_entries = []
            except (ValueError, TypeError):
                spray_entries = []

            try:
                report_findings = _json.loads(request.POST.get('report_findings_json') or '[]')
                if not isinstance(report_findings, list):
                    report_findings = []
            except (ValueError, TypeError):
                report_findings = []

            _SIG_PREFIX = 'data:image/png;base64,'
            if client_sig and not client_sig.startswith(_SIG_PREFIX):
                client_sig = ''
            if supervisor_sig and not supervisor_sig.startswith(_SIG_PREFIX):
                supervisor_sig = ''

            def _to_int(val):
                try:
                    v = int(val)
                    return v if v >= 0 else None
                except (ValueError, TypeError):
                    return None

            order.status              = 'completed'
            order.workers_count       = _to_int(workers_raw)
            order.vehicles_count      = _to_int(vehicles_raw)
            order.building_type       = building_type
            order.spray_entries       = spray_entries
            order.report_findings     = report_findings
            order.supervisor_notes    = sup_notes
            order.report_submitted_by = request.user
            order.report_submitted_at = timezone.now()
            if not order.close_date:
                order.close_date = timezone.now().date()
            if client_sig:
                order.client_signature = client_sig
            if supervisor_sig:
                order.supervisor_signature = supervisor_sig
            order.save(update_fields=[
                'status', 'workers_count', 'vehicles_count',
                'building_type', 'spray_entries', 'report_findings',
                'supervisor_notes', 'report_submitted_by', 'report_submitted_at',
                'close_date', 'client_signature', 'supervisor_signature',
            ])
            report_photos = request.FILES.getlist('report_photos')
            for photo in report_photos:
                FieldWorkPhoto.objects.create(
                    work_order=order, phase='work',
                    file=photo, uploaded_by=request.user,
                )
            photos_all = order.photos.order_by('uploaded_at')
            _fw_log(order, 'completed', request.user)
            success = 'تم حفظ تقرير المراقب — الحالة: تم إنجاز الخدمة.'

        elif action == 'postpone_order' and can_submit_report:
            from datetime import date as _date
            postponed_until_raw = (request.POST.get('postponed_until') or '').strip()
            postpone_notes      = (request.POST.get('postpone_notes') or '').strip()
            try:
                postponed_until = _date.fromisoformat(postponed_until_raw)
                if postponed_until <= _date.today():
                    errors.append('يرجى اختيار تاريخ مستقبلي.')
            except ValueError:
                postponed_until = None
                errors.append('يرجى تحديد تاريخ التأجيل.')
            if not errors:
                order.status              = 'postponed_client'
                order.postponed_until     = postponed_until
                order.assigned_supervisor = None
                order.assigned_at         = None
                if postpone_notes:
                    order.supervisor_notes = postpone_notes
                order.save(update_fields=[
                    'status', 'postponed_until', 'supervisor_notes',
                    'assigned_supervisor', 'assigned_at',
                ])
                _fw_log(order, 'postponed', request.user,
                        to_value=postponed_until.strftime('%d/%m/%Y'),
                        note=postpone_notes)
                success = f'تم تأجيل الموعد إلى {postponed_until.strftime("%d/%m/%Y")}.'

        # ── Close request ───────────────────────────────────────────────────
        elif action == 'close_request' and (can_submit_report or can_edit):
            close_reason = (request.POST.get('close_reason') or '').strip()
            proof = request.FILES.get('closure_proof')
            valid_close_reasons = {
                'closed_private_building', 'closed_no_answer', 'closed_other_municipal',
                'closed_observation', 'closed_low_infestation', 'closed_moderate_infestation',
                'closed_high_infestation', 'closed_out_of_service', 'closed_customer_refused',
                'closed_mobile_off', 'closed_not_attending', 'closed_not_available',
                'closed_scheduled_client',
            }
            if order.status in _FW_CLOSED_STATUSES:
                errors.append('الطلب مغلق بالفعل.')
            elif close_reason not in valid_close_reasons:
                errors.append('يرجى اختيار سبب الإغلاق.')
            elif not proof:
                errors.append('يرجى رفع صورة إثبات.')
            else:
                ext = os.path.splitext(proof.name)[1].lower()
                if ext not in ALLOWED_PHOTO_EXTENSIONS:
                    errors.append('يُسمح فقط بصور JPG أو PNG.')
                else:
                    order.status = close_reason
                    order.close_date = timezone.now().date()
                    order.no_answer_screenshot = proof
                    order.report_submitted_by = request.user
                    order.report_submitted_at = timezone.now()
                    order.save(update_fields=[
                        'status', 'close_date', 'no_answer_screenshot',
                        'report_submitted_by', 'report_submitted_at',
                    ])
                    _fw_log(order, 'closed', request.user,
                            to_value=_STATUS_LABELS.get(close_reason, close_reason))
                    success = 'تم إغلاق الطلب بنجاح.'

        # ── Admin quick close (private company / other municipality) ─────────
        elif action == 'admin_quick_close' and can_admin:
            quick_close_reason = (request.POST.get('quick_close_reason') or '').strip()
            valid_quick_close_reasons = {
                'private_company', 'other_municipal',
                'closed_out_of_service', 'closed_customer_refused',
                'closed_not_available', 'closed_not_attending',
                'closed_private_building', 'closed_other_municipal',
            }
            if order.status in _FW_CLOSED_STATUSES:
                errors.append('الطلب مغلق بالفعل.')
            elif quick_close_reason not in valid_quick_close_reasons:
                errors.append('يرجى اختيار سبب الإغلاق.')
            else:
                order.status = quick_close_reason
                order.close_date = timezone.now().date()
                order.report_submitted_by = request.user
                order.report_submitted_at = timezone.now()
                order.save(update_fields=[
                    'status', 'close_date', 'report_submitted_by', 'report_submitted_at',
                ])
                _fw_log(order, 'closed', request.user,
                        to_value=_STATUS_LABELS.get(quick_close_reason, quick_close_reason))
                success = 'تم إغلاق الطلب بنجاح.'

        # ── Reopen closed order ──────────────────────────────────────────────
        elif action == 'reopen_order' and can_admin:
            if order.status not in _FW_CLOSED_STATUSES:
                errors.append('الطلب ليس مغلقاً.')
            else:
                old_status_label = _STATUS_LABELS.get(order.status, order.status)
                order.status             = 'new'
                order.close_date         = None
                order.report_submitted_at = None
                order.report_submitted_by = None
                order.save(update_fields=[
                    'status', 'close_date', 'report_submitted_at', 'report_submitted_by',
                ])
                _fw_log(order, 'status_changed', request.user,
                        from_value=old_status_label, to_value='إعادة فتح')
                success = 'تم إعادة فتح الطلب — الحالة: قيد الانتظار.'

        # ── Delete photo ─────────────────────────────────────────────────────
        elif action == 'delete_photo' and can_edit:
            photo_id = (request.POST.get('photo_id') or '').strip()
            try:
                photo = FieldWorkPhoto.objects.get(id=photo_id, work_order=order)
                photo.file.delete(save=False)
                photo.delete()
                success = 'تم حذف الصورة.'
                photos_all = order.photos.order_by('uploaded_at')
            except FieldWorkPhoto.DoesNotExist:
                errors.append('الصورة غير موجودة.')

        if not errors:
            return redirect('field_work_detail', pk=order.pk)

    fw_supervisor_users = _fw_supervisor_users_qs() if can_assign else []
    is_closed = order.status in _FW_CLOSED_STATUSES
    can_close = (can_submit_report or can_edit) and not is_closed
    from datetime import date as _date
    _BUILDING_TYPE_CHOICES = [
        'Villa', 'Government Facility', 'School', 'Public Park',
        'Mosque', 'Palace', 'Labor Accommodation', 'Vip Villa',
    ]
    return render(request, 'hcsd/field_work_detail.html', {
        'order': order,
        'can_edit': can_edit,
        'can_admin': can_admin,
        'can_assign': can_assign,
        'can_submit_report': can_submit_report,
        'can_receive': can_receive,
        'can_close': can_close,
        'is_closed': is_closed,
        'fw_supervisor_users': fw_supervisor_users,
        'photos_all': photos_all,
        'errors': errors,
        'success': success,
        'building_type_choices': _BUILDING_TYPE_CHOICES,
        'today_date': _date.today().isoformat(),
        'order_logs': order.logs.all(),
    })


# ---------------------------------------------------------------------------
# Print Report
# ---------------------------------------------------------------------------

_ACTIVE_INGREDIENTS = {
    "Actellic 50 EC": "Primiphos Methyl 50%",
    "ECOLARVACIDE EC": "Temephos 50%",
    "Starycide SC 480": "Triflumuron 48%",
    "DIFRON 25 SC": "Diflubenzuron 25%",
    "GRAYBATE 50 SG": "Temephos 50 g/kg",
    "BIOPREN 4 GR": "S-Methoprene 0.4%",
    "LAROXYFEN PLUS WT": "Pyriproxyfen 2%",
    "Aqua k-Othrine": "Deltamethrin 20 g/L",
    "Chirotox": "Tetramethrin 5% / Piperonyl Butoxide 5%",
    "TETRACON 50 EC": "Lambda-cyhalothrin 50 g/L / Tetramethrin 10 g/L / PBO 20 g/L",
    "KULCYPERIN 100/3 EC": "Cypermethrin 10% / Tetramethrin 0.5% / Piperonyl Butoxide 2.5%",
    "DEMON MAX INSECTICIDE": "Cypermethrin 25.3% (w/w)",
    "Solfac EC 50": "Cyfluthrin 5%",
    "CYMPERATOR 25 EC": "Cypermethrin 26%",
    "Bio Amplat": "Cypermethrin 93% min. / Tetramethrin 94% min. / Piperonyl Butoxide 94% min.",
    "ROTRYN 200": "Cypermethrin 20% w/w (200 g/L)",
    "GUADIN SE": "Dinotefuran 10%",
    "BAITFURAN SP": "Dinotefuran 12%",
    "Detral Super": "Deltamethrin 0.7 g / Esbiothrin 0.7 g / Piperonyl Butoxide 7 g",
    "K-Othrine Partix": "Deltamethrin 25 g/L",
    "PERMETHOR": "Permethrin 10 g/kg",
    "Temprid SC": "Imidacloprid 21% / Beta-Cyfluthrin 10.5%",
    "HYMENOPHTHOR GR": "Fipronil 0.1 g/kg",
    "Vertox Oktablok": "Brodifacoum 0.005%",
    "FACORAT PELLETS": "Brodifacoum 0.005 g",
    "VICTOR V FAST-KILL BRAND BLOCKS II": "Bromethalin 0.01%",
    "SUREFIRE ALL WEATHER BLOCKS": "Brodifacoum 0.05 g/kg",
    "PROTECT SENSATION 2IN1": "Bromadiolone 0.005% (0.05 g/kg)",
    "VERTOX PASTA BAIT": "Brodifacoum 0.005% (0.05 g/kg)",
    "STELLIOX D50": "Difenacoum 0.005% W/W",
    "TALON WB": "Brodifacoum 0.005%",
    "SUREFIRE BROMA BLOCKS RODENTICIDE": "Bromadiolone 0.05 g/kg",
    "NOCURAT PARAFFINATO": "Difenacoum",
    "BuyBlocker Snake Deter": "Cedar Oil 1.00% / Cinnamon Oil 0.60% / Clove Oil 0.40%",
    "BOOM": "Clove Oil / Peppermint Oil / Citronella Oil / Cedar Oil / Cinnamon Oil / Thyme Oil",
    "CYPFORCE 40 EC": "Cypermethrin 25% / Tetramethrin 5% / Piperonyl Butoxide 10%",
    "D-TETRASUPER EC": "D-Tetramethrin 10%",
    "TEMEPHOS 55EC": "Temephos 50%",
    "Petrol": "Petroleum Distillates",
    "Diesel": "Diesel Fuel",
}

_DANGEROUS_PESTS = frozenset(['Bees', 'Scorpion', 'Snake', 'Poisonous Spider', 'Wasp And Bee', 'Other Harmful Pest'])
_NUISANCE_PESTS  = frozenset(['Ants', 'German Cockroach', 'American Cockroach', 'Lizard', 'Termites', 'Non Poisonous Spider', 'Drain Flies', 'Fruit Flies'])
_VECTOR_PESTS    = frozenset(['Mosquito Adult', 'Mosquito Aedes', 'Mosquito Culex', 'Mosquito Anopheles', 'Rodents Roof Rat', 'Rodents Norway Rat', 'Rodents House mouse', 'House Flies'])
_OTHERS_PESTS    = frozenset(['Agricultural Pest'])

# ── Monthly Excel export helpers ─────────────────────────────────────────────
_MONTH_ABBR = ['JAN', 'FEB', 'MAR', 'APR', 'MAY', 'JUN',
                'JUL', 'AUG', 'SEP', 'OCT', 'NOV', 'DEC']



def _user_display_name(user):
    if not user:
        return ''
    name = user.get_full_name()
    return name if name else user.username


def _sup_name(user) -> str:
    """Return the supervisor's Arabic profile name, falling back to full name / username."""
    if not user:
        return ''
    try:
        name_ar = user.fw_supervisor_profile.name_ar
        if name_ar:
            return name_ar
    except Exception:
        pass
    return _user_display_name(user)




def _pest_category(pests):
    cats = []
    if any(p in _DANGEROUS_PESTS for p in pests):
        cats.append('Dangerous')
    if any(p in _NUISANCE_PESTS for p in pests):
        cats.append('Nuisance')
    if any(p in _VECTOR_PESTS for p in pests):
        cats.append('Vector')
    if any(p in _OTHERS_PESTS for p in pests):
        cats.append('Others')
    return ' / '.join(cats) if cats else ''


@login_required
def field_work_report_print(request, pk):
    order = get_object_or_404(
        FieldWorkOrder.objects.select_related('report_submitted_by'),
        pk=pk,
    )
    materials = []
    observations = []
    all_pests = []
    seen_pests = set()
    for entry in (order.spray_entries or []):
        loc = entry.get('location', '')
        pests = entry.get('pests', [])
        pesticides = entry.get('pesticides', [])
        for p in pesticides:
            materials.append({
                'product': p.get('name', ''),
                'qty': f"{p.get('qty', '')} {p.get('unit', '')}".strip(),
                'active_ingredient': _ACTIVE_INGREDIENTS.get(p.get('name', ''), ''),
                'location': loc,
            })
        for p in pests:
            if p not in seen_pests:
                seen_pests.add(p)
                all_pests.append(p)
        infestation_raw = entry.get('infestation', '')
        infestation_label = _INFESTATION_LABELS.get(infestation_raw, infestation_raw)
        entry_pesticides = [
            f"{p.get('name', '')} {p.get('qty', '')} {p.get('unit', '')}".strip()
            for p in pesticides if p.get('name')
        ]
        observations.append({
            'building_type': order.building_type or '',
            'observation': infestation_label,
            'pest_category': _pest_category(pests),
            'pest_found': ', '.join(pests),
            'location': loc,
            'action': ', '.join(entry.get('actions', [])) or entry.get('action', ''),
            'pesticides': entry_pesticides,
        })
    time_in_dt = order.time_in or order.location_saved_at
    close_date = order.close_date or (
        timezone.localtime(order.report_submitted_at).date()
        if order.report_submitted_at else None
    )
    photos = list(order.photos.order_by('uploaded_at'))
    is_rescheduled   = bool(order.postponed_until)
    is_completed     = bool(order.report_submitted_at)
    is_force_closed  = order.status.startswith('closed_') and is_completed
    proof_img_url    = order.no_answer_screenshot.url if order.no_answer_screenshot else None
    return render(request, 'hcsd/field_work_report_print.html', {
        'order': order,
        'materials': materials,
        'observations': observations,
        'all_pests': all_pests,
        'report_findings': order.report_findings or [],
        'time_in_dt': time_in_dt,
        'close_date': close_date,
        'photos': photos,
        'is_rescheduled': is_rescheduled,
        'is_completed': is_completed,
        'is_force_closed': is_force_closed,
        'proof_img_url': proof_img_url,
    })


# ---------------------------------------------------------------------------
# Excel Report (per-order)
# ---------------------------------------------------------------------------

_INFESTATION_LABELS = {
    'no_infestation':       'No Infestation',
    'low_infestation':      'Low Infestation',
    'moderate_infestation': 'Moderate Infestation',
    'high_infestation':     'High Infestation',
}


@login_required
def field_work_excel_report(request, pk):
    order = get_object_or_404(
        FieldWorkOrder.objects.select_related(
            'created_by', 'assigned_supervisor', 'received_by',
            'report_submitted_by', 'location_saved_by',
        ),
        pk=pk,
    )

    try:
        import openpyxl
        from openpyxl.styles import (
            Alignment, Border, Font, PatternFill, Side,
        )
        from openpyxl.utils import get_column_letter
    except ImportError:
        return HttpResponse('openpyxl is not installed.', status=500)

    wb = openpyxl.Workbook()

    # ── Shared styles ────────────────────────────────────────────────────────
    def _thin():
        s = Side(style='thin', color='AAAAAA')
        return Border(left=s, right=s, top=s, bottom=s)

    FILL_HEADER   = PatternFill('solid', fgColor='1A5276')
    FILL_SUBHDR   = PatternFill('solid', fgColor='2471A3')
    FILL_LABEL    = PatternFill('solid', fgColor='EBF5FB')
    FILL_SECTION  = PatternFill('solid', fgColor='D6EAF8')

    FONT_HDR      = Font(name='Calibri', bold=True, color='FFFFFF', size=11)
    FONT_SUBHDR   = Font(name='Calibri', bold=True, color='FFFFFF', size=10)
    FONT_LABEL    = Font(name='Calibri', bold=True, size=10)
    FONT_VAL      = Font(name='Calibri', size=10)
    FONT_SECTION  = Font(name='Calibri', bold=True, size=10, color='1A5276')

    CENTER = Alignment(horizontal='center', vertical='center', wrap_text=True)
    LEFT   = Alignment(horizontal='left',   vertical='center', wrap_text=True)

    def _hdr_cell(ws, row, col, text, width_hint=None):
        c = ws.cell(row=row, column=col, value=text)
        c.font   = FONT_HDR
        c.fill   = FILL_HEADER
        c.border = _thin()
        c.alignment = CENTER
        if width_hint:
            ws.column_dimensions[get_column_letter(col)].width = width_hint
        return c

    def _subhdr_cell(ws, row, col, text):
        c = ws.cell(row=row, column=col, value=text)
        c.font   = FONT_SUBHDR
        c.fill   = FILL_SUBHDR
        c.border = _thin()
        c.alignment = CENTER
        return c

    def _lbl(ws, row, col, text):
        c = ws.cell(row=row, column=col, value=text)
        c.font   = FONT_LABEL
        c.fill   = FILL_LABEL
        c.border = _thin()
        c.alignment = LEFT
        return c

    def _val(ws, row, col, text):
        c = ws.cell(row=row, column=col, value=text)
        c.font   = FONT_VAL
        c.border = _thin()
        c.alignment = LEFT
        return c

    def _section(ws, row, col, text, span_end_col=None):
        c = ws.cell(row=row, column=col, value=text)
        c.font   = FONT_SECTION
        c.fill   = FILL_SECTION
        c.border = _thin()
        c.alignment = LEFT
        if span_end_col and span_end_col > col:
            ws.merge_cells(
                start_row=row, start_column=col,
                end_row=row, end_column=span_end_col,
            )
        return c

    def _dt_str(dt):
        if not dt:
            return ''
        local = timezone.localtime(dt)
        return local.strftime('%d/%m/%Y %H:%M')

    def _d_str(d):
        if not d:
            return ''
        return d.strftime('%d/%m/%Y')

    def _user_str(u):
        if not u:
            return ''
        name = u.get_full_name()
        return name if name else u.username

    # ════════════════════════════════════════════════════════════════════════
    # Sheet 1 — Order Summary
    # ════════════════════════════════════════════════════════════════════════
    ws1 = wb.active
    ws1.title = 'Order Summary'
    ws1.sheet_view.rightToLeft = False

    # Title row
    ws1.merge_cells('A1:D1')
    title_cell = ws1['A1']
    title_cell.value = f'Field Work Order Report  —  #{order.order_number or order.pk}'
    title_cell.font  = Font(name='Calibri', bold=True, size=14, color='1A5276')
    title_cell.alignment = CENTER
    title_cell.fill = PatternFill('solid', fgColor='D6EAF8')
    ws1.row_dimensions[1].height = 30

    ws1.column_dimensions['A'].width = 26
    ws1.column_dimensions['B'].width = 32
    ws1.column_dimensions['C'].width = 26
    ws1.column_dimensions['D'].width = 32

    row = 2

    def _pair(label1, val1, label2='', val2=''):
        nonlocal row
        _lbl(ws1, row, 1, label1)
        _val(ws1, row, 2, val1)
        _lbl(ws1, row, 3, label2)
        _val(ws1, row, 4, val2)
        ws1.row_dimensions[row].height = 16
        row += 1

    def _full_row(label, value):
        nonlocal row
        _lbl(ws1, row, 1, label)
        c = _val(ws1, row, 2, value)
        ws1.merge_cells(
            start_row=row, start_column=2,
            end_row=row, end_column=4,
        )
        ws1.row_dimensions[row].height = 16
        row += 1

    def _sec(title):
        nonlocal row
        _section(ws1, row, 1, title, span_end_col=4)
        ws1.row_dimensions[row].height = 18
        row += 1

    # General info
    _sec('General Information')
    _pair('Order Number',   order.order_number or str(order.pk),
          'Status',         order.get_status_display())
    _pair('Customer Name',  order.customer_name or order.site_name or '',
          'Mobile',         order.mobile or '')
    _pair('Area',           order.area or order.location or '',
          'Street / House', f"{order.street_number or ''} / {order.house_number or ''}".strip(' /'))
    _pair('Complaint Source', order.get_complaint_source_display() or '—',
          'Building Type',     order.building_type or '')
    _pair('Work Type',      order.work_type or '', '', '')
    _pair('Workers',        order.workers_count if order.workers_count is not None else '',
          'Vehicles',       order.vehicles_count if order.vehicles_count is not None else '')
    if order.description:
        _full_row('Description', order.description)
    if order.notes:
        _full_row('Notes', order.notes)
    if order.supervisor_notes:
        _full_row('Supervisor Notes', order.supervisor_notes)

    row += 1  # blank

    # Timeline
    _sec('Timeline / Lifecycle')
    _pair('Created At',        _dt_str(order.created_at),
          'Created By',        _user_str(order.created_by))
    _pair('Assigned At',       _dt_str(order.assigned_at),
          'Assigned Supervisor', _user_str(order.assigned_supervisor))
    _pair('Received At',       _dt_str(order.received_at),
          'Received By',       _user_str(order.received_by))
    _pair('GPS Saved At',      _dt_str(order.location_saved_at),
          'GPS Saved By',      _user_str(order.location_saved_by))
    if order.gps_lat is not None:
        _pair('GPS Latitude',  order.gps_lat,
              'GPS Longitude', order.gps_lng)
    _pair('Time In',           _dt_str(order.time_in or order.location_saved_at), '', '')
    _pair('Report Submitted',  _dt_str(order.report_submitted_at),
          'Submitted By',      _user_str(order.report_submitted_by))
    close_date = order.close_date or (
        timezone.localtime(order.report_submitted_at).date()
        if order.report_submitted_at else None
    )
    _pair('Close Date',        _d_str(close_date), '', '')
    if order.postponed_until:
        _pair('Postponed Until', _d_str(order.postponed_until), '', '')

    # ════════════════════════════════════════════════════════════════════════
    # Sheet 2 — Spray Entries
    # ════════════════════════════════════════════════════════════════════════
    ws2 = wb.create_sheet('Spray Entries')
    ws2.sheet_view.rightToLeft = False

    hdr2_cols = [
        ('Location', 22),
        ('Infestation', 18),
        ('Pests Found', 30),
        ('Pest Category', 16),
        ('Actions Taken', 35),
        ('Findings', 35),
    ]
    for col_idx, (hdr, w) in enumerate(hdr2_cols, start=1):
        _hdr_cell(ws2, 1, col_idx, hdr, width_hint=w)
    ws2.row_dimensions[1].height = 20

    spray_entries = order.spray_entries or []
    for r_idx, entry in enumerate(spray_entries, start=2):
        pests = entry.get('pests', [])
        infest_key = entry.get('infestation', '')
        infest_label = _INFESTATION_LABELS.get(infest_key, infest_key)
        actions  = entry.get('actions', []) or ([entry.get('action')] if entry.get('action') else [])
        findings = entry.get('findings', [])

        row_vals = [
            entry.get('location', ''),
            infest_label,
            ', '.join(pests),
            _pest_category(pests),
            '\n'.join(actions),
            '\n'.join(findings),
        ]
        for col_idx, val in enumerate(row_vals, start=1):
            c = ws2.cell(row=r_idx, column=col_idx, value=val)
            c.font      = FONT_VAL
            c.border    = _thin()
            c.alignment = Alignment(
                horizontal='left', vertical='top', wrap_text=True,
            )
        ws2.row_dimensions[r_idx].height = max(
            15 * max(len(actions), len(findings), 1), 16,
        )

    if not spray_entries:
        ws2.cell(row=2, column=1, value='— No spray entries recorded —').font = Font(
            name='Calibri', italic=True, color='999999',
        )

    # ════════════════════════════════════════════════════════════════════════
    # Sheet 3 — Materials Used
    # ════════════════════════════════════════════════════════════════════════
    ws3 = wb.create_sheet('Materials Used')
    ws3.sheet_view.rightToLeft = False

    hdr3_cols = [
        ('Location', 22),
        ('Product', 30),
        ('Qty', 10),
        ('Unit', 12),
        ('Active Ingredient', 45),
    ]
    for col_idx, (hdr, w) in enumerate(hdr3_cols, start=1):
        _hdr_cell(ws3, 1, col_idx, hdr, width_hint=w)
    ws3.row_dimensions[1].height = 20

    mat_row = 2
    for entry in spray_entries:
        loc = entry.get('location', '')
        for p in entry.get('pesticides', []):
            name = p.get('name', '')
            vals = [
                loc,
                name,
                p.get('qty', ''),
                p.get('unit', ''),
                _ACTIVE_INGREDIENTS.get(name, ''),
            ]
            for col_idx, val in enumerate(vals, start=1):
                c = ws3.cell(row=mat_row, column=col_idx, value=val)
                c.font      = FONT_VAL
                c.border    = _thin()
                c.alignment = LEFT
            ws3.row_dimensions[mat_row].height = 16
            mat_row += 1

    if mat_row == 2:
        ws3.cell(row=2, column=1, value='— No materials recorded —').font = Font(
            name='Calibri', italic=True, color='999999',
        )

    # ── Output ───────────────────────────────────────────────────────────────
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    order_ref = order.order_number or str(order.pk)
    filename  = f'field_work_{order_ref}.xlsx'
    response  = HttpResponse(
        buf.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Monthly Excel Report (all orders, one sheet per month, matching template)
# ---------------------------------------------------------------------------

_MONTHLY_HEADERS = [
    'رقم الطلب', 'تاريخ الطلب', 'تاريخ الإغلاق', 'اسم المتعامل',
    'حالة الطلب', 'سبب الإغلاق', 'تاريخ التأجيل',
    'الموبايل', 'رقم الشارع', 'رقم المنزل',
    'المنطقة', 'نوع الحشرات', 'المراقب المسؤول', 'مصدر الشكوى',
]

_MONTHLY_COL_WIDTHS = [
    12, 12, 12, 25,
    16, 34, 14,
    14, 10, 10,
    22, 32, 25, 16,
]

# Maps status → short category label for col 6
_STATUS_CATEGORY = {
    'completed':        'مكتمل',
    'postponed_client': 'مؤجل',
    'other_municipal':  'بلدية أخرى',
    'new':              'جديد',
    'assigned':         'معين',
    'received':         'تم الاستلام',
}

def _status_category(status: str) -> str:
    if status in _STATUS_CATEGORY:
        return _STATUS_CATEGORY[status]
    if status.startswith('closed_'):
        return 'مغلق'
    return status

# Maps status → detailed closure reason for col 7 (only for closed/postponed)
_CLOSURE_REASON = {
    'completed':                    '',
    'postponed_client':             'تأجيل من العميل',
    'other_municipal':              'تابع لبلدية أخرى',
    'closed_private_building':      'شركة نظافة خاصة (داخل بناية)',
    'closed_no_answer':             'لم يرد العميل على الهاتف',
    'closed_other_municipal':       'تابع لبلدية أخرى',
    'closed_observation':           'ملاحظة',
    'closed_low_infestation':       'تفشٍ خفيف',
    'closed_moderate_infestation':  'تفشٍ متوسط',
    'closed_high_infestation':      'تفشٍ شديد',
    'closed_out_of_service':        'خارج نطاق الخدمة',
    'closed_customer_refused':      'العميل رفض الخدمة',
    'closed_mobile_off':            'هاتف العميل مغلق',
    'closed_not_attending':         'العميل لا يرد على المكالمات',
    'closed_not_available':         'العميل غير متاح',
    'closed_scheduled_client':      'تم الجدولة من قِبل العميل',
}


@login_required
def field_work_monthly_excel(request):
    if not (_can_admin(request.user) or _can_data_entry(request.user)):
        return redirect('field_work_list')

    from collections import defaultdict

    try:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return HttpResponse('openpyxl is not installed.', status=500)

    today = timezone.localtime(timezone.now()).date()
    raw_from = (request.GET.get('date_from') or '').strip()
    raw_to   = (request.GET.get('date_to')   or '').strip()

    try:
        date_from = _dt.date.fromisoformat(raw_from)
    except ValueError:
        date_from = today.replace(month=1, day=1)

    try:
        date_to = _dt.date.fromisoformat(raw_to)
    except ValueError:
        date_to = today

    if date_from > date_to:
        date_from, date_to = date_to, date_from

    from django.db.models import Q
    status_filters = request.GET.getlist('status_filter')
    include_all = 'all' in status_filters or not status_filters

    status_q = Q()
    if not include_all:
        if 'completed' in status_filters:
            status_q |= Q(status='completed')
        if 'closed' in status_filters:
            status_q |= Q(status__startswith='closed_')
        if 'postponed' in status_filters:
            status_q |= Q(status='postponed_client')

    date_q = (
        Q(request_date__gte=date_from, request_date__lte=date_to) |
        Q(request_date__isnull=True, created_at__date__gte=date_from, created_at__date__lte=date_to)
    )

    orders = (
        FieldWorkOrder.objects
        .filter(date_q)
        .filter(status_q if not include_all else Q())
        .select_related(
            'assigned_supervisor', 'assigned_supervisor__fw_supervisor_profile',
            'received_by', 'received_by__fw_supervisor_profile',
        )
        .order_by('request_date', 'created_at', 'pk')
    )

    monthly = defaultdict(list)
    for order in orders:
        group_date = order.request_date or timezone.localtime(order.created_at).date()
        monthly[(group_date.year, group_date.month)].append(order)

    wb = openpyxl.Workbook()
    wb.remove(wb.active)

    _thin = Side(style='thin', color='AAAAAA')
    _border = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)

    FILL_INFO  = PatternFill('solid', fgColor='1F4E79')
    FONT_WHITE = Font(name='Calibri', bold=True, color='FFFFFF', size=9)
    FONT_DATA  = Font(name='Calibri', size=9)
    ALIGN_CTR  = Alignment(horizontal='center', vertical='center', wrap_text=False)
    ALIGN_LFT  = Alignment(horizontal='left',   vertical='center', wrap_text=False)

    for (yr, month) in sorted(monthly.keys()):
        orders_in_month = monthly[(yr, month)]
        if not orders_in_month:
            continue

        yr2 = str(yr)[-2:]
        ws = wb.create_sheet(title=f'{_MONTH_ABBR[month - 1]} - {yr2}')
        ws.sheet_view.rightToLeft = True
        ws.freeze_panes = 'A2'

        for col_idx, width in enumerate(_MONTHLY_COL_WIDTHS, start=1):
            ws.column_dimensions[get_column_letter(col_idx)].width = width

        ws.row_dimensions[1].height = 30
        for col_idx, hdr in enumerate(_MONTHLY_HEADERS, start=1):
            c = ws.cell(row=1, column=col_idx, value=hdr)
            c.border    = _border
            c.alignment = ALIGN_CTR
            c.font      = FONT_WHITE
            c.fill = FILL_INFO

        for r_idx, order in enumerate(orders_in_month, start=2):
            ws.row_dimensions[r_idx].height = 20

            sup_name = (
                _sup_name(order.received_by)
                or _sup_name(order.assigned_supervisor)
                or order.supervisor_name or ''
            )

            # Collect unique pests from spray_entries; fall back to pest_types text
            seen_pests: set = set()
            ordered_pests: list = []
            for entry in (order.spray_entries or []):
                for p in entry.get('pests', []):
                    if p not in seen_pests:
                        seen_pests.add(p)
                        ordered_pests.append(p)
            pest_col = order.pest_types or ', '.join(ordered_pests)

            row_vals = [
                order.order_number or '',                             # 1
                order.request_date,                                   # 2
                order.close_date,                                     # 3
                order.customer_name or order.site_name or '',         # 4
                _status_category(order.status),                       # 5 حالة الطلب
                _CLOSURE_REASON.get(order.status, ''),                # 6 سبب الإغلاق
                order.postponed_until,                                 # 7 تاريخ التأجيل
                order.mobile or '',                                   # 8
                order.street_number or '',                            # 9
                order.house_number or '',                             # 10
                order.area or order.location or '',                   # 11
                pest_col,                                             # 12 نوع الحشرات
                sup_name,                                             # 13
                order.get_complaint_source_display() or '',           # 14 مصدر الشكوى
            ]

            for col_idx, val in enumerate(row_vals, start=1):
                c = ws.cell(row=r_idx, column=col_idx, value=val)
                c.border    = _border
                c.font      = FONT_DATA
                if col_idx in (2, 3, 7, 8, 9, 10):
                    c.alignment = ALIGN_CTR
                else:
                    c.alignment = ALIGN_LFT

    if not wb.sheetnames:
        ws = wb.create_sheet('No Data')
        ws.cell(row=1, column=1, value=f'No orders with close date between {date_from} and {date_to}.')

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    filename = f'FieldWork_{date_from}_to_{date_to}.xlsx'
    response = HttpResponse(
        buf.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Monthly Materials Excel (one row per pesticide per order)
# ---------------------------------------------------------------------------

_MAT_HEADERS = [
    'رقم الطلب', 'تاريخ الطلب', 'تاريخ التنفيذ', 'تاريخ الإغلاق',
    'اسم المتعامل', 'حالة الطلب', 'الموبايل',
    'رقم الشارع', 'رقم المنزل', 'المنطقة', 'نوع المبنى',
    'المراقب', 'نوع الحشرات', 'مصدر الشكوى',
    'اسم المادة', 'المادة الفعالة', 'الكمية', 'الوحدة',
    'مكان التطبيق', 'الآفات المستهدفة', 'فئة الآفة',
]

_MAT_COL_WIDTHS = [
    12, 12, 12, 12,
    28, 22, 14,
    10, 10, 22, 20,
    25, 28, 16,
    22, 22, 10, 8,
    22, 30, 18,
]


@login_required
def field_work_materials_excel(request):
    if not (_can_admin(request.user) or _can_data_entry(request.user)):
        return redirect('field_work_list')

    try:
        import openpyxl
        from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        return HttpResponse('openpyxl is not installed.', status=500)

    today = timezone.localtime(timezone.now()).date()
    raw_from = (request.GET.get('date_from') or '').strip()
    raw_to   = (request.GET.get('date_to')   or '').strip()

    try:
        date_from = _dt.date.fromisoformat(raw_from)
    except ValueError:
        date_from = today.replace(month=1, day=1)

    try:
        date_to = _dt.date.fromisoformat(raw_to)
    except ValueError:
        date_to = today

    if date_from > date_to:
        date_from, date_to = date_to, date_from

    from django.db.models import Q
    status_filters = request.GET.getlist('status_filter')
    include_all = 'all' in status_filters or not status_filters

    status_q = Q()
    if not include_all:
        if 'completed' in status_filters:
            status_q |= Q(status='completed')
        if 'closed' in status_filters:
            status_q |= Q(status__startswith='closed_')
        if 'postponed' in status_filters:
            status_q |= Q(status='postponed_client')

    date_q = (
        Q(request_date__gte=date_from, request_date__lte=date_to) |
        Q(request_date__isnull=True, created_at__date__gte=date_from, created_at__date__lte=date_to)
    )

    orders = (
        FieldWorkOrder.objects
        .filter(date_q)
        .filter(status_q if not include_all else Q())
        .select_related(
            'assigned_supervisor', 'assigned_supervisor__fw_supervisor_profile',
            'received_by', 'received_by__fw_supervisor_profile',
        )
        .order_by('request_date', 'created_at', 'pk')
    )

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f'{date_from} to {date_to}'
    ws.sheet_view.rightToLeft = True
    ws.freeze_panes = 'A2'

    _thin   = Side(style='thin', color='AAAAAA')
    _border = Border(left=_thin, right=_thin, top=_thin, bottom=_thin)

    FILL_HDR  = PatternFill('solid', fgColor='1F4E79')
    FONT_HDR  = Font(name='Calibri', bold=True, color='FFFFFF', size=9)
    FONT_DATA = Font(name='Calibri', size=9)
    ALIGN_CTR = Alignment(horizontal='center', vertical='center', wrap_text=False)
    ALIGN_LFT = Alignment(horizontal='right',  vertical='center', wrap_text=False)

    for col_idx, width in enumerate(_MAT_COL_WIDTHS, start=1):
        ws.column_dimensions[get_column_letter(col_idx)].width = width

    ws.row_dimensions[1].height = 30
    for col_idx, hdr in enumerate(_MAT_HEADERS, start=1):
        c = ws.cell(row=1, column=col_idx, value=hdr)
        c.fill      = FILL_HDR
        c.font      = FONT_HDR
        c.border    = _border
        c.alignment = ALIGN_CTR

    r = 1
    for order in orders:
        sup_name = (
            _sup_name(order.received_by)
            or _sup_name(order.assigned_supervisor)
            or order.supervisor_name or ''
        )

        base = [
            order.order_number or '',                                # 1
            order.request_date,                                      # 2
            order.work_date,                                         # 3
            order.close_date,                                        # 4
            order.customer_name or order.site_name or '',            # 5
            order.get_status_display(),                              # 6
            order.mobile or '',                                      # 7
            order.street_number or '',                               # 8
            order.house_number or '',                                # 9
            order.area or order.location or '',                      # 10
            order.building_type or '',                               # 11
            sup_name,                                                # 12
            order.pest_types or '',                                  # 13
            order.get_complaint_source_display() or '',              # 14 مصدر الشكوى
        ]

        # Build material rows from spray_entries
        mat_rows = []
        for entry in (order.spray_entries or []):
            loc        = entry.get('location', '')
            pests      = entry.get('pests', [])
            pest_cat   = _pest_category(pests)
            pests_str  = ', '.join(pests)
            pesticides = entry.get('pesticides', [])
            for p in pesticides:
                name = p.get('name', '').strip()
                if not name:
                    continue
                mat_rows.append([
                    name,                                            # 15 اسم المادة
                    _ACTIVE_INGREDIENTS.get(name, ''),               # 16 المادة الفعالة
                    p.get('qty', ''),                                # 17 الكمية
                    p.get('unit', ''),                               # 18 الوحدة
                    loc,                                             # 19 مكان التطبيق
                    pests_str,                                       # 20 الآفات المستهدفة
                    pest_cat,                                        # 21 فئة الآفة
                ])

        if not mat_rows:
            mat_rows = [['', '', '', '', '', '', '']]

        for mat in mat_rows:
            r += 1
            ws.row_dimensions[r].height = 15
            row_vals = base + mat
            for col_idx, val in enumerate(row_vals, start=1):
                c = ws.cell(row=r, column=col_idx, value=val)
                c.border    = _border
                c.font      = FONT_DATA
                c.alignment = ALIGN_CTR if col_idx in (2, 3, 4, 7, 8, 9, 17, 18) else ALIGN_LFT

    if r == 1:
        ws.cell(row=2, column=1, value=f'No orders between {date_from} and {date_to}.')

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)

    filename = f'FieldWork_Materials_{date_from}_to_{date_to}.xlsx'
    response = HttpResponse(
        buf.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Word Report
# ---------------------------------------------------------------------------

@login_required
def field_work_report(request, pk):
    order = get_object_or_404(FieldWorkOrder.objects.select_related('created_by'), pk=pk)

    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return HttpResponse('مكتبة python-docx غير مثبتة.', status=500)

    doc = Document()

    # ── Page setup: A4, Arabic RTL ──────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = int(8.27 * 914400)
    section.page_height = int(11.69 * 914400)
    section.left_margin = section.right_margin = int(1 * 914400)
    section.top_margin  = section.bottom_margin = int(1 * 914400)

    def set_rtl(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def add_heading(text, level=1):
        p = doc.add_paragraph()
        set_rtl(p)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(16 if level == 1 else 13)
        run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
        return p

    def add_field(label, value):
        p = doc.add_paragraph()
        set_rtl(p)
        lbl = p.add_run(f'{label}: ')
        lbl.bold = True
        lbl.font.size = Pt(11)
        val = p.add_run(str(value) if value else '—')
        val.font.size = Pt(11)

    def add_section_title(text):
        p = doc.add_paragraph()
        set_rtl(p)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x2a, 0x6a, 0x9a)
        doc.add_paragraph()

    # ── Title ───────────────────────────────────────────────────────────────
    title = doc.add_paragraph()
    set_rtl(title)
    t = title.add_run('تقرير أمر العمل الميداني')
    t.bold = True
    t.font.size = Pt(20)
    t.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    doc.add_paragraph()

    # ── Basic info ──────────────────────────────────────────────────────────
    add_heading('معلومات الطلب', level=1)
    add_field('رقم الأمر', f'#{order.pk}')
    add_field('نوع العمل', order.work_type)
    add_field('اسم الموقع', order.site_name or None)
    add_field('العنوان', order.location or None)
    add_field('تاريخ التنفيذ', order.work_date.strftime('%d/%m/%Y') if order.work_date else None)
    add_field('تاريخ الإنشاء', order.created_at.strftime('%d/%m/%Y'))
    add_field('الحالة', order.get_status_display())
    doc.add_paragraph()

    # ── Work details ────────────────────────────────────────────────────────
    add_heading('تفاصيل العمل', level=1)
    add_field('عدد العمال', order.workers_count)
    add_field('المعدات المستخدمة', order.equipment_used or None)
    add_field('اكتملت العملية', 'نعم' if order.work_completed is True else ('لا' if order.work_completed is False else None))
    if order.description:
        add_field('وصف العمل', order.description)
    if order.notes:
        add_field('ملاحظات', order.notes)
    doc.add_paragraph()

    # ── Photos ──────────────────────────────────────────────────────────────
    # Photos are re-encoded to a capped resolution/quality before embedding —
    # original phone-camera photos can be several MB each, which made the
    # generated .docx huge and slow to build/download for orders with many photos.
    from PIL import Image, ImageOps
    _PHOTO_MAX_DIM = 1000
    _PHOTO_JPEG_QUALITY = 70

    all_photos = list(order.photos.order_by('uploaded_at'))
    if all_photos:
        add_section_title('صور العمل')
        for photo in all_photos:
            try:
                photo_path = photo.file.path
                if not os.path.exists(photo_path):
                    continue
                with Image.open(photo_path) as img:
                    img = ImageOps.exif_transpose(img).convert('RGB')
                    img.thumbnail((_PHOTO_MAX_DIM, _PHOTO_MAX_DIM), Image.LANCZOS)
                    photo_buf = io.BytesIO()
                    img.save(photo_buf, format='JPEG', quality=_PHOTO_JPEG_QUALITY, optimize=True)
                    photo_buf.seek(0)
                doc.add_picture(photo_buf, width=Inches(4.5))
                if photo.caption:
                    cap = doc.add_paragraph(photo.caption)
                    set_rtl(cap)
                doc.add_paragraph()
            except Exception:
                logger.exception('Failed to add photo %s to Word report', photo.pk)

    # ── Output ──────────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f'field_work_{order.pk}.docx'
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


# ---------------------------------------------------------------------------
# Excel Import
# ---------------------------------------------------------------------------

_EXCEL_SESSION_KEY    = 'fw_excel_import'
_MAX_EXCEL_ROWS       = 2000
_ALLOWED_EXCEL_EXTS   = {'.xlsx', '.xls'}

_COL_MAP = {
    'رقم الطلب': 'order_number',    'رقم الامر': 'order_number',
    'رقم الأمر': 'order_number',    'order no': 'order_number',
    'order number': 'order_number', 'order_no': 'order_number',
    'تاريخ الطلب': 'request_date',  'request date': 'request_date',
    'تاريخ الإغلاق': 'close_date',  'تاريخ الاغلاق': 'close_date',
    'close date': 'close_date',
    'اسم المتعامل': 'customer_name','اسم العميل': 'customer_name',
    'المتعامل': 'customer_name',    'customer': 'customer_name',
    'customer name': 'customer_name',
    'الموبايل': 'mobile',           'الجوال': 'mobile',
    'رقم الهاتف': 'mobile',         'mobile': 'mobile',
    'phone': 'mobile',
    'المنطقة': 'area',              'area': 'area',
    'رقم الشارع': 'street_number',  'الشارع': 'street_number',
    'street': 'street_number',      'street no': 'street_number',
    'رقم المنزل': 'house_number',   'house': 'house_number',
    'house no': 'house_number',
    'نوع الحشرات': 'pest_types',    'الحشرات': 'pest_types',
    'pest': 'pest_types',           'pest type': 'pest_types',
    'المشرف المعالج': 'supervisor_name', 'المشرف': 'supervisor_name',
    'supervisor': 'supervisor_name',
    'العامل': 'worker_name',        'worker': 'worker_name',
    'حالة الطلب': 'excel_status',   'الحالة': 'excel_status',
    'status': 'excel_status',
    'ملاحظة الحالة': 'excel_status_note', 'ملاحظة': 'excel_status_note',
    'note': 'excel_status_note',    'notes': 'excel_status_note',
    'الشهر': 'month_sheet',         'month': 'month_sheet',
}


def _norm_header(h):
    return ' '.join(str(h or '').strip().lower().split())


def _parse_xl_date(val):
    if val is None:
        return ''
    if isinstance(val, (_dt.datetime, _dt.date)):
        d = val.date() if isinstance(val, _dt.datetime) else val
        return d.isoformat()
    s = str(val).strip()
    for fmt in ('%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y', '%m/%d/%Y'):
        try:
            return _dt.datetime.strptime(s, fmt).date().isoformat()
        except ValueError:
            pass
    return s


def _to_date(s):
    if not s:
        return None
    try:
        return _dt.date.fromisoformat(s)
    except ValueError:
        return None


def _str_val(v):
    return '' if v is None else str(v).strip()


def _extract_excel_rows(file_obj):
    """Return (rows, error_message). rows = list[dict]."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(file_obj, read_only=True, data_only=True)
    except Exception:
        return [], 'تعذّر قراءة الملف — تأكد أنه ملف Excel صالح (.xlsx).'

    ws = wb.active
    it = ws.iter_rows(values_only=True)

    header_row = None
    for row in it:
        if any(c is not None for c in row):
            header_row = row
            break
    if header_row is None:
        return [], 'الملف فارغ أو لا يحتوي على بيانات.'

    col_map = {}
    for idx, cell in enumerate(header_row):
        field = _COL_MAP.get(_norm_header(cell))
        if field:
            col_map[idx] = field

    if not col_map:
        return [], 'لم يتم التعرف على أعمدة الملف. تأكد أن الصف الأول يحتوي على عناوين الأعمدة.'

    rows = []
    for row in it:
        if not any(c is not None for c in row):
            continue
        rec = {}
        for idx, field in col_map.items():
            val = row[idx] if idx < len(row) else None
            rec[field] = _parse_xl_date(val) if field in ('request_date', 'close_date') else _str_val(val)
        if not any(v for v in rec.values()):
            continue
        rows.append(rec)
        if len(rows) >= _MAX_EXCEL_ROWS:
            break

    return rows, ''


@login_required
def field_work_excel_import(request):
    if not (_can_admin(request.user) or _can_data_entry(request.user)):
        return redirect('field_work_list')

    error = ''
    if request.method == 'POST':
        f = request.FILES.get('excel_file')
        if not f:
            error = 'يرجى اختيار ملف Excel.'
        elif os.path.splitext(f.name)[1].lower() not in _ALLOWED_EXCEL_EXTS:
            error = 'يُسمح فقط بملفات .xlsx أو .xls'
        else:
            rows, err = _extract_excel_rows(f)
            if err:
                error = err
            elif not rows:
                error = 'لم يتم العثور على صفوف بيانات في الملف.'
            else:
                request.session[_EXCEL_SESSION_KEY] = rows
                return redirect('field_work_excel_review')

    return render(request, 'hcsd/field_work_excel_import.html', {'error': error})


@login_required
def field_work_excel_review(request):
    if not (_can_admin(request.user) or _can_data_entry(request.user)):
        return redirect('field_work_list')

    rows = request.session.get(_EXCEL_SESSION_KEY)
    if not rows:
        return redirect('field_work_excel_import')

    order_numbers = {r.get('order_number', '') for r in rows if r.get('order_number')}
    existing = set(
        FieldWorkOrder.objects.filter(order_number__in=order_numbers)
        .values_list('order_number', flat=True)
    )

    enriched = [{**r, 'is_dup': r.get('order_number', '') in existing} for r in rows]
    new_count = sum(1 for r in enriched if not r['is_dup'])
    dup_count = len(rows) - new_count

    if request.method == 'POST':
        mode = request.POST.get('import_mode', 'new_only')
        row_count = len(rows)
        to_create = []
        for i in range(row_count):
            include = request.POST.get(f'row_{i}_include')
            if not include:
                continue
            order_number = (request.POST.get(f'row_{i}_order_number') or '').strip()
            if mode == 'new_only' and order_number in existing:
                continue
            _ex_status = (request.POST.get(f'row_{i}_excel_status') or '').strip()
            to_create.append(FieldWorkOrder(
                order_number    = order_number,
                request_date    = _to_date((request.POST.get(f'row_{i}_request_date') or '').strip()),
                customer_name   = (request.POST.get(f'row_{i}_customer_name') or '').strip(),
                mobile          = (request.POST.get(f'row_{i}_mobile') or '').strip(),
                area            = (request.POST.get(f'row_{i}_area') or '').strip(),
                house_number    = (request.POST.get(f'row_{i}_house_number') or '').strip(),
                pest_types      = (request.POST.get(f'row_{i}_pest_types') or '').strip(),
                supervisor_name = (request.POST.get(f'row_{i}_supervisor_name') or '').strip(),
                worker_name     = (request.POST.get(f'row_{i}_worker_name') or '').strip(),
                excel_status    = _ex_status,
                status          = _infer_status_from_excel(_ex_status),
                street_number   = rows[i].get('street_number', ''),
                close_date      = _to_date(rows[i].get('close_date', '')),
                excel_status_note = rows[i].get('excel_status_note', ''),
                month_sheet     = rows[i].get('month_sheet', ''),
                source            = 'excel',
                complaint_source  = 'hotline',
                created_by        = request.user,
            ))
        FieldWorkOrder.objects.bulk_create(to_create, batch_size=500)
        created = len(to_create)
        del request.session[_EXCEL_SESSION_KEY]
        return redirect(reverse('field_work_list') + f'?imported={created}')

    return render(request, 'hcsd/field_work_excel_review.html', {
        'rows':      enriched,
        'total':     len(rows),
        'new_count': new_count,
        'dup_count': dup_count,
    })


# ---------------------------------------------------------------------------
# Supervisors Management
# ---------------------------------------------------------------------------

@login_required
def field_work_supervisors(request):
    if not (_can_admin(request.user) or _can_data_entry(request.user)):
        return redirect('field_work_list')

    from django.contrib.auth.models import User
    from django.db.models import Count, Q

    supervisors = (
        _fw_supervisor_users_qs()
        .prefetch_related('fw_supervisor_areas', 'fw_supervisor_profile')
        .annotate(
            active_count=Count(
                'field_work_assigned',
                filter=~Q(field_work_assigned__status__in=_FW_CLOSED_STATUSES),
                distinct=True,
            ) + Count(
                'field_work_received',
                filter=~Q(field_work_received__status__in=_FW_CLOSED_STATUSES),
                distinct=True,
            ),
        )
    )

    # Collect all distinct areas from existing orders for the area dropdown
    existing_areas = (
        FieldWorkOrder.objects.exclude(area='')
        .values_list('area', flat=True)
        .distinct()
        .order_by('area')
    )

    error = ''
    success = ''

    if request.method == 'POST':
        if not _can_admin(request.user):
            return redirect('field_work_supervisors')

        action = request.POST.get('action', '')

        if action == 'create_supervisor':
            from django.contrib.auth.models import Group
            username     = (request.POST.get('username') or '').strip()
            password     = (request.POST.get('password') or '').strip()
            password2    = (request.POST.get('password2') or '').strip()
            name_ar      = (request.POST.get('name_ar') or '').strip()
            name_en      = (request.POST.get('name_en') or '').strip()
            admin_number = (request.POST.get('admin_number') or '').strip()
            areas_raw    = (request.POST.get('areas') or '').strip()

            if not username:
                error = 'يرجى إدخال اسم المستخدم.'
            elif User.objects.filter(username=username).exists():
                error = 'اسم المستخدم مستخدم بالفعل.'
            elif not password:
                error = 'يرجى إدخال كلمة المرور.'
            elif password != password2:
                error = 'كلمتا المرور غير متطابقتان.'
            else:
                user = User.objects.create_user(
                    username=username, password=password,
                    first_name=name_ar, last_name=name_en,
                )
                FieldWorkSupervisorProfile.objects.create(
                    user=user, name_ar=name_ar, name_en=name_en,
                    admin_number=admin_number,
                )
                grp, _ = Group.objects.get_or_create(name='fw_supervisor')
                user.groups.add(grp)
                for area in [a.strip() for a in areas_raw.split('\n') if a.strip()]:
                    FieldWorkSupervisorArea.objects.get_or_create(
                        supervisor=user, area=area,
                        defaults={'assigned_by': request.user},
                    )
                success = f'تم إنشاء حساب المراقب "{name_ar or username}" بنجاح.'

        elif action == 'add_area':
            sup_id = (request.POST.get('supervisor_id') or '').strip()
            area = (request.POST.get('area') or '').strip()
            if not sup_id or not area:
                error = 'يرجى اختيار مراقب ومنطقة.'
            else:
                try:
                    sup = User.objects.get(pk=int(sup_id), groups__name__in=['fw_supervisor', 'Field Work Supervisor'])
                    _, created = FieldWorkSupervisorArea.objects.get_or_create(
                        supervisor=sup, area=area,
                        defaults={'assigned_by': request.user},
                    )
                    success = f'تمت إضافة منطقة "{area}" للمراقب {sup.get_full_name() or sup.username}.' if created else 'المنطقة مضافة مسبقاً.'
                except (ValueError, User.DoesNotExist):
                    error = 'المراقب غير صالح.'
            if not error:
                from django.urls import reverse
                return redirect(reverse('field_work_supervisors') + f'?sel_sup={sup_id}')

        elif action == 'remove_area':
            area_id = (request.POST.get('area_id') or '').strip()
            try:
                FieldWorkSupervisorArea.objects.filter(pk=int(area_id)).delete()
                success = 'تم حذف المنطقة.'
            except (ValueError, Exception):
                error = 'تعذّر الحذف.'

        elif action == 'bulk_reassign' and _can_admin(request.user):
            from django.utils import timezone as _tz
            order_ids_raw = request.POST.getlist('order_ids')
            new_sup_id    = (request.POST.get('new_supervisor_id') or '').strip()
            try:
                order_ids = [int(i) for i in order_ids_raw if i.strip()]
                new_sup   = _fw_supervisor_users_qs().get(pk=int(new_sup_id))
                if order_ids:
                    FieldWorkOrder.objects.filter(pk__in=order_ids).update(
                        assigned_supervisor=new_sup,
                        assigned_at=_tz.now(),
                        received_by=None,
                        received_at=None,
                        status='supervisor_assigned',
                    )
                    name = new_sup.get_full_name() or new_sup.username
                    success = f'تم إعادة تعيين {len(order_ids)} أمر للمراقب {name}.'
                else:
                    error = 'لم يتم تحديد أي أوامر.'
            except (ValueError, Exception):
                error = 'حدث خطأ أثناء إعادة التعيين.'

        if not error:
            return redirect('field_work_supervisors')

    sel_sup = (request.GET.get('sel_sup') or '').strip()
    try:
        sel_sup_id = int(sel_sup)
    except (ValueError, TypeError):
        sel_sup_id = None

    # Build active-orders map for the drawer panel
    from collections import defaultdict
    _status_labels = dict(FieldWorkOrder.STATUS_CHOICES)

    _active = list(
        FieldWorkOrder.objects.filter(~Q(status__in=_FW_CLOSED_STATUSES))
        .values('id', 'order_number', 'customer_name', 'site_name',
                'area', 'status', 'assigned_supervisor_id', 'received_by_id')
    )
    _sup_orders = defaultdict(list)
    _seen = set()
    for o in _active:
        entry = {
            'id':     o['id'],
            'num':    o['order_number'] or str(o['id']),
            'name':   o['customer_name'] or o['site_name'] or '—',
            'area':   o['area'] or '—',
            'status': _status_labels.get(o['status'], o['status']),
        }
        for sid in {o['assigned_supervisor_id'], o['received_by_id']}:
            if sid and (sid, o['id']) not in _seen:
                _seen.add((sid, o['id']))
                _sup_orders[sid].append(entry)

    return render(request, 'hcsd/field_work_supervisors.html', {
        'supervisors':          supervisors,
        'existing_areas':       existing_areas,
        'fw_supervisor_users':  _fw_supervisor_users_qs(),
        'error':                error,
        'success':              success,
        'sel_sup_id':           sel_sup_id,
        'sup_orders_json':      dict(_sup_orders),
    })


# ---------------------------------------------------------------------------
# Supervisor order history
# ---------------------------------------------------------------------------

@login_required
def field_work_supervisor_orders(request, pk):
    from django.core.paginator import Paginator
    from django.db.models import Q, Count
    from django.contrib.auth.models import User

    is_admin      = _can_admin(request.user)
    is_data_entry = _can_data_entry(request.user)
    is_self       = request.user.pk == pk

    if not (is_admin or is_data_entry or is_self):
        return redirect('field_work_list')

    supervisor = get_object_or_404(
        User.objects.prefetch_related('fw_supervisor_profile'), pk=pk
    )

    status_filter = (request.GET.get('status') or 'all').strip()
    search_q      = (request.GET.get('q')       or '').strip()
    quick_filter  = (request.GET.get('quick')   or '').strip()
    reassigned    = request.GET.get('reassigned')

    # POST: bulk reassign active orders to another supervisor
    if request.method == 'POST' and (is_admin or is_data_entry):
        from django.utils import timezone as _tz
        order_ids_raw = request.POST.getlist('order_ids')
        new_sup_id    = (request.POST.get('new_supervisor_id') or '').strip()
        sf  = (request.POST.get('status_filter') or 'all').strip()
        q   = (request.POST.get('q') or '').strip()
        try:
            order_ids = [int(i) for i in order_ids_raw if i.strip()]
            new_sup   = _fw_supervisor_users_qs().get(pk=int(new_sup_id))
            updated   = (
                FieldWorkOrder.objects
                .filter(pk__in=order_ids)
                .exclude(status__in=_FW_CLOSED_STATUSES)
                .update(
                    assigned_supervisor=new_sup,
                    assigned_at=_tz.now(),
                    received_by=None,
                    received_at=None,
                    status='supervisor_assigned',
                )
            )
        except Exception:
            updated = 0
        from urllib.parse import urlencode
        params = urlencode({k: v for k, v in {
            'status': sf, 'q': q, 'reassigned': updated,
        }.items() if v})
        return redirect(
            reverse('field_work_supervisor_orders', kwargs={'pk': pk}) + f'?{params}'
        )

    qs = (
        FieldWorkOrder.objects
        .filter(Q(assigned_supervisor_id=pk) | Q(received_by_id=pk))
        .only(
            'id', 'order_number', 'customer_name', 'site_name',
            'area', 'mobile', 'status', 'request_date', 'work_date',
            'assigned_at', 'received_at', 'created_at', 'source',
        )
        .order_by('-created_at', '-pk')
    )

    if quick_filter == 'today':
        today = timezone.localdate()
        qs = qs.filter(Q(request_date=today) | Q(request_date__isnull=True, created_at__date=today))
    elif quick_filter == 'new':
        qs = qs.filter(status__in=['new', 'supervisor_assigned'])
    elif quick_filter == 'received':
        qs = qs.filter(status__in=['supervisor_assigned', 'order_received'])
    elif quick_filter == 'completed':
        qs = qs.filter(status='completed')
    elif quick_filter == 'closed':
        qs = qs.filter(status__in=_FW_TRULY_CLOSED)
    elif quick_filter == 'postponed':
        qs = qs.filter(status='postponed_client')
    elif status_filter == 'closed':
        qs = qs.filter(status__in=_FW_TRULY_CLOSED)
    elif status_filter != 'all':
        qs = qs.filter(status=status_filter)
    if search_q:
        from django.db.models import Q as _Q
        qs = qs.filter(
            _Q(order_number__icontains=search_q)
            | _Q(mobile__icontains=search_q)
            | _Q(area__icontains=search_q)
        )

    counts = (
        FieldWorkOrder.objects
        .filter(Q(assigned_supervisor_id=pk) | Q(received_by_id=pk))
        .values('status').annotate(n=Count('id'))
    )
    status_counts = {c['status']: c['n'] for c in counts}
    total        = sum(status_counts.values())
    active_count = sum(v for s, v in status_counts.items() if s not in _FW_CLOSED_STATUSES)
    closed_count = total - active_count

    paginator = Paginator(qs, 30)
    page_obj  = paginator.get_page(request.GET.get('page', 1))

    try:
        profile = supervisor.fw_supervisor_profile
    except Exception:
        profile = None

    return render(request, 'hcsd/field_work_supervisor_orders.html', {
        'supervisor':        supervisor,
        'profile':           profile,
        'page_obj':          page_obj,
        'status_filter':          status_filter,
        'active_status_choices':  [(v, l) for v, l in FieldWorkOrder.STATUS_CHOICES if v not in _FW_TRULY_CLOSED],
        'status_labels_ar':       dict(FieldWorkOrder.STATUS_CHOICES),
        'status_labels_en':  FieldWorkOrder.STATUS_LABELS_EN,
        'closed_statuses':   list(_FW_CLOSED_STATUSES),
        'fw_supervisors':    _fw_supervisor_users_qs().prefetch_related('fw_supervisor_profile'),
        'total':             total,
        'active_count':      active_count,
        'closed_count':      closed_count,
        'can_admin':         is_admin or is_data_entry,
        'reassigned':  reassigned,
        'search_q':    search_q,
        'quick_filter': quick_filter,
    })

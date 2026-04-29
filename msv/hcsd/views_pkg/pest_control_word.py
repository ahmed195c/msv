import io

import qrcode
from django.conf import settings
from django.contrib.auth.decorators import login_required
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from django.urls import reverse
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..models import PirmetClearance
from .common import _activities_for_enginer, _restricted_activities_for_enginer

ACTIVITY_LABELS_AR = {
    'public_health_pest_control': 'مكافحة آفات الصحة العامة',
    'flying_insects': 'مكافحة الحشرات الطائرة',
    'rodents': 'مكافحة القوارض',
    'termite_control': 'مكافحة النمل الأبيض',
    'grain_pests': 'مكافحة آفات الحبوب',
}

AR_CONDITIONS = [
    ('الإشتراطات:', [
        'يجب الاحتفاظ بسجلات تداول مبيدات المكافحة.',
        'يمنع تجارة أو بيع المبيدات على الأفراد أو المؤسسات.',
        'غير مصرح باستخدام المواد المستحلبة في الأماكن المغلقة.',
        'يسمح بإستخدام المبيدات المصرحة والمسجلة لدى وزارة التغير المناخي والبيئة فقط.',
        'في حال غياب المهندس دون إخطار البلدية برسالة رسمية فإن ذلك يعرض الشركة للإجراءات القانونية السارية في البلدية.',
        'يسري هذا التصريح على ممارسة نشاط مكافحة الآفات الصحة العامة داخل حدود مدينة الشارقة.',
    ]),
    ('بنود التصريح:', [
        'يتم إبراز التصريح الأصلي لموظف البلدية عند الطلب.',
        'الإخلال بالإشتراطات المذكورة يعرضك للإجراءات القانونية السارية في البلدية.',
        'يتم تجديد أو إلغاء هذا التصريح خلال مدة أسبوع قبل تاريخ انتهاء التصريح.',
        'للبلدية الحق في إيقاف هذا التصريح في حال الاخلال بالاشتراطات المعمول بها أو إلغاءه عند الضرورة.',
    ]),
]

EN_CONDITIONS = [
    ('Recommendations:', [
        'All documents of pesticides handling must be kept.',
        'Pesticides trading or pesticides sailing exhibit for individuals and facilities.',
        'Emulsion concentrate (EC) in closed places is not allow.',
        'Only approved pesticide from Ministry of Climate Change and Environment are allow to use.',
        'In case of Engineer leave on vacation without inform Waste Control and regulating section in Sharjah City Municipality that exposed company to fine.',
        'This permit is valid within Sharjah City only.',
    ]),
    ('Permit Requirements', [
        'Original of Public health pest control must be showing to Sharjah City Municipality members if needed.',
        "Company exposed to fine if not follow all the recommendations of permit mention above.",
        'Public health pest control permit must be renewing or cancelling within one week before the date of expire.',
        "Sharjah City Municipality has a right to suspend Public health pest control permit when the company don't follow the requirements also has a right to cancel if needed.",
    ]),
]


# ── XML helpers ──────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _bidi_para(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def _bidi_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)


def _styled_run(para, text, bold=False, size_pt=12, color='111111', rtl=True):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor.from_string(color)
    if rtl:
        _bidi_run(run)
    return run


def _cell_write(cell, text, bold=False, size_pt=12, color='111111',
                align=WD_ALIGN_PARAGRAPH.CENTER, rtl=True):
    para = cell.paragraphs[0]
    para.alignment = align
    if rtl:
        _bidi_para(para)
    _styled_run(para, text, bold=bold, size_pt=size_pt, color=color, rtl=rtl)


def _add_para(doc, text='', bold=False, size_pt=12, color='111111',
              align=WD_ALIGN_PARAGRAPH.CENTER, rtl=True):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_after = Pt(3)
    if rtl:
        _bidi_para(para)
    if text:
        _styled_run(para, text, bold=bold, size_pt=size_pt, color=color, rtl=rtl)
    return para


# ── Table building helpers ────────────────────────────────────────────────────

def _section_header(table, ar_text, en_text='', cols=3):
    row = table.add_row()
    cell = row.cells[0]
    if cols > 1:
        for i in range(1, cols):
            cell.merge(row.cells[i])
    _set_cell_bg(cell, '3d3d3d')
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bidi_para(para)
    run = para.add_run(ar_text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _bidi_run(run)
    if en_text:
        run2 = para.add_run(f'  / {en_text}')
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def _data_row(table, ar_label, value, en_label, val_size=12):
    row = table.add_row()
    c0, c1, c2 = row.cells
    _set_cell_bg(c0, 'dce8db')
    _cell_write(c0, ar_label, bold=True, size_pt=12, color='173317', rtl=True)
    _cell_write(c1, value or '—', bold=True, size_pt=val_size, color='111111', rtl=True)
    _set_cell_bg(c2, 'eef4ed')
    _cell_write(c2, en_label, bold=False, size_pt=10, color='3a5a3a', rtl=False)


def _add_title_table(doc, ar_title, en_title):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    _set_cell_bg(cell, '3d3d3d')
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _bidi_para(para)
    r1 = para.add_run(ar_title)
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    _bidi_run(r1)
    r2 = para.add_run(f'\n{en_title}')
    r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    return tbl


def _make_qr_stream(url_path):
    full_url = settings.SITE_URL.rstrip('/') + url_path
    img = qrcode.make(full_url, box_size=6, border=2)
    buf = io.BytesIO()
    img.save(buf, format='PNG')
    buf.seek(0)
    return buf


def _add_dept_header(doc):
    for line in ['إدارة الرقابة والسلامة الصحية', 'قسم رقابة وتنظيم النفايات']:
        p = _add_para(doc, line, bold=True, size_pt=13, color='1a2e1a',
                      align=WD_ALIGN_PARAGRAPH.RIGHT, rtl=True)
        p.paragraph_format.space_after = Pt(2)


def _add_cond_section(cell, heading, items, rtl=True):
    align = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
    ph = cell.add_paragraph()
    ph.alignment = align
    if rtl:
        _bidi_para(ph)
    rh = ph.add_run(heading)
    rh.bold = True
    rh.font.size = Pt(13)
    if rtl:
        _bidi_run(rh)

    for i, item in enumerate(items, 1):
        pi = cell.add_paragraph()
        pi.alignment = align
        if rtl:
            _bidi_para(pi)
        ri = pi.add_run(f'{i}. {item}')
        ri.font.size = Pt(11)
        if rtl:
            _bidi_run(ri)

    cell.add_paragraph()


# ── Main view ─────────────────────────────────────────────────────────────────

@login_required
def pest_control_permit_word(request, id):
    pirmet = get_object_or_404(
        PirmetClearance.objects.select_related('company', 'company__enginer'),
        id=id,
        permit_type='pest_control',
    )
    company = pirmet.company
    enginer = company.enginer if company else None
    allowed = _activities_for_enginer(enginer)
    restricted = _restricted_activities_for_enginer(enginer)

    doc = Document()

    # Page setup — A4
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = Cm(1.5)
    sec.right_margin = Cm(1.5)
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)

    # ── PAGE 1 ────────────────────────────────────────────────────────────────

    _add_dept_header(doc)

    _add_title_table(
        doc,
        'تصريح مزاولة نشاط مكافحة آفات الصحة العامة',
        'Public Health Pest Control Activity Permit',
    )

    _add_para(doc)  # spacer

    # Permit number + dates
    tbl_permit = doc.add_table(rows=0, cols=3)
    tbl_permit.style = 'Table Grid'

    row = tbl_permit.add_row()
    _set_cell_bg(row.cells[0], 'dce8db')
    _cell_write(row.cells[0], 'رقم التصريح', bold=True, size_pt=12, color='173317', rtl=True)
    _cell_write(row.cells[1], pirmet.permit_no or '—', bold=True, size_pt=17, color='111111', rtl=True)
    _set_cell_bg(row.cells[2], 'eef4ed')
    _cell_write(row.cells[2], 'Permit No.', size_pt=10, color='3a5a3a', rtl=False)

    _data_row(tbl_permit, 'تاريخ الإصدار',
              pirmet.issue_date.strftime('%d/%m/%Y') if pirmet.issue_date else '—', 'Issue Date')
    _data_row(tbl_permit, 'تاريخ الانتهاء',
              pirmet.dateOfExpiry.strftime('%d/%m/%Y') if pirmet.dateOfExpiry else '—', 'Expiry Date')

    # Company
    tbl_company = doc.add_table(rows=0, cols=3)
    tbl_company.style = 'Table Grid'
    _section_header(tbl_company, 'بيانات شركة مكافحة آفات الصحة العامة',
                    'Public Health Pest Control Company')
    _data_row(tbl_company, 'اسم الشركة', company.name if company else '—', 'Company Name')
    _data_row(tbl_company, 'رقم الرخصة التجارية', company.number if company else '—', 'Trade License No.')
    _data_row(tbl_company, 'تاريخ انتهاء الرخصة',
              company.trade_license_exp.strftime('%d/%m/%Y') if (company and company.trade_license_exp) else '—',
              'License Expiry')
    _data_row(tbl_company, 'عنوان الشركة', company.address if company else '—', 'Address')
    _data_row(tbl_company, 'رقم التواصل',
              (company.owner_phone or company.landline or '—') if company else '—', 'Contact No.')
    _data_row(tbl_company, 'البريد الإلكتروني', company.email if company else '—', 'Email')

    # Engineer
    tbl_eng = doc.add_table(rows=0, cols=3)
    tbl_eng.style = 'Table Grid'
    _section_header(tbl_eng, 'بيانات المهندس المسؤول', 'Responsible Engineer')
    if enginer:
        _data_row(tbl_eng, 'اسم المهندس', enginer.name or '—', 'Engineer Name')
        _data_row(tbl_eng, 'رقم الهاتف', enginer.phone or '—', 'Contact No.')
    else:
        row = tbl_eng.add_row()
        cell = row.cells[0]
        cell.merge(row.cells[1]).merge(row.cells[2])
        _cell_write(cell, 'لا يوجد مهندس مسجل / No engineer registered', size_pt=11, color='777777')

    # Activities
    tbl_act = doc.add_table(rows=0, cols=2)
    tbl_act.style = 'Table Grid'
    _section_header(tbl_act, 'بيانات الأنشطة', 'Activity Details', cols=2)

    ah = tbl_act.add_row()
    _set_cell_bg(ah.cells[0], 'dce8db')
    _cell_write(ah.cells[0], 'الأنشطة المصرح بها', bold=True, size_pt=12, color='173317', rtl=True)
    _set_cell_bg(ah.cells[1], 'f5dada')
    _cell_write(ah.cells[1], 'الأنشطة غير المصرح بها', bold=True, size_pt=12, color='7a2a2a', rtl=True)

    ac = tbl_act.add_row()
    allowed_cell = ac.cells[0]
    restricted_cell = ac.cells[1]

    first_allowed = True
    for item in allowed:
        label = f'• {ACTIVITY_LABELS_AR.get(item, item)}'
        if first_allowed:
            p = allowed_cell.paragraphs[0]
            first_allowed = False
        else:
            p = allowed_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bidi_para(p)
        r = p.add_run(label)
        r.font.size = Pt(12)
        _bidi_run(r)
    if pirmet.allowed_other:
        p = allowed_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _bidi_para(p)
        r = p.add_run(f'• أخرى: {pirmet.allowed_other}')
        r.font.size = Pt(12)
        _bidi_run(r)

    first_restricted = True
    if restricted or pirmet.restricted_other:
        for item in restricted:
            label = f'• {ACTIVITY_LABELS_AR.get(item, item)}'
            if first_restricted:
                p = restricted_cell.paragraphs[0]
                first_restricted = False
            else:
                p = restricted_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _bidi_para(p)
            r = p.add_run(label)
            r.font.size = Pt(12)
            _bidi_run(r)
        if pirmet.restricted_other:
            p = restricted_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _bidi_para(p)
            r = p.add_run(f'• أخرى: {pirmet.restricted_other}')
            r.font.size = Pt(12)
            _bidi_run(r)
    else:
        _cell_write(restricted_cell, '—', size_pt=11, color='999999')

    # Payment
    tbl_pay = doc.add_table(rows=0, cols=3)
    tbl_pay.style = 'Table Grid'
    _data_row(tbl_pay, 'رقم ايصال الدفع', pirmet.PaymentNumber or '—', 'Receipt Voucher No.')
    _data_row(tbl_pay, 'تاريخ الايصال',
              pirmet.payment_date.strftime('%d/%m/%Y') if pirmet.payment_date else '—', 'Receipt Date')

    # Footer: QR + stamp (2-col table)
    _add_para(doc)
    tbl_footer = doc.add_table(rows=1, cols=2)
    tbl_footer.style = 'Table Grid'
    qr_cell = tbl_footer.rows[0].cells[0]
    stamp_cell = tbl_footer.rows[0].cells[1]

    qr_stream = _make_qr_stream(reverse('pest_control_permit_print', args=[pirmet.id]))
    qr_para = qr_cell.paragraphs[0]
    qr_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    qr_para.add_run().add_picture(qr_stream, width=Cm(2.5))

    _cell_write(stamp_cell, 'قسم رقابة وتنظيم النفايات',
                bold=True, size_pt=12, color='1a3a1a', rtl=True)
    stamp_cell.add_paragraph()
    p2 = stamp_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Waste Control and Regulation Section')
    r2.font.size = Pt(10)
    r2.bold = True
    r2.font.color.rgb = RGBColor(0x1a, 0x3a, 0x1a)

    # ── PAGE 2 — Conditions ────────────────────────────────────────────────────

    doc.add_page_break()

    _add_dept_header(doc)

    tbl_cond = doc.add_table(rows=1, cols=2)
    tbl_cond.style = 'Table Grid'
    ar_cell = tbl_cond.rows[0].cells[0]
    en_cell = tbl_cond.rows[0].cells[1]

    # Remove default empty first paragraph in each cell before adding content
    for section_title, items in AR_CONDITIONS:
        _add_cond_section(ar_cell, section_title, items, rtl=True)

    for section_title, items in EN_CONDITIONS:
        _add_cond_section(en_cell, section_title, items, rtl=False)

    # ── Serialize ──────────────────────────────────────────────────────────────

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = (pirmet.permit_no or str(pirmet.id)).replace('/', '-')
    filename = f'permit_{safe_name}.docx'
    response = HttpResponse(
        buf.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response
